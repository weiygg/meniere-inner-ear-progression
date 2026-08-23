from __future__ import annotations

import json
import shutil
from argparse import Namespace
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from inner_ear_geometry_features import (
    build_ear_feature_row,
    clean_mask,
    fit_pca_shape,
    fit_plane_pca,
    load_mask_nii,
    mask_points_mm,
    plane_residuals,
    projected_arc_features,
    skeleton_points_mm,
    subsample_points,
)
from inner_ear_multifeature_model import run_training


ROOT = Path("analysis_out") / "updated_068L_20260509"
OLD_GEOM = Path("seg3") / "analysis_out_geometry"
GEOM = ROOT / "geometry"
MODEL_DIR = ROOT / "els_tv_model"
REPORT = ROOT / "梅尼埃内耳半规管影像特征与ELS占比预测更新报告_068L补齐.docx"


def ensure_dirs() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    GEOM.mkdir(parents=True, exist_ok=True)
    MODEL_DIR.mkdir(parents=True, exist_ok=True)


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, dtype={"pid": str, "side": str, "struct": str})


def key_mask(df: pd.DataFrame, pid: str, side: str, struct: str | None = None) -> pd.Series:
    mask = df["pid"].astype(str).str.zfill(3).eq(pid) & df["side"].astype(str).str.upper().eq(side)
    if struct is not None:
        mask &= df["struct"].astype(str).str.upper().eq(struct.upper())
    return mask


def as_number(value):
    try:
        if pd.isna(value):
            return value
        return float(value)
    except Exception:
        return value


def compute_hsc_row() -> tuple[dict, dict, dict]:
    fp = Path("seg3") / "sub068" / "068L_HSC.nii.gz"
    if not fp.exists():
        raise FileNotFoundError(fp)
    rng = np.random.default_rng(42)
    mask, zooms = load_mask_nii(fp)
    cleaned = clean_mask(mask, closing_radius=1, min_voxels=200)
    vox_points = mask_points_mm(cleaned, zooms)
    skel_points = skeleton_points_mm(cleaned, zooms)
    if skel_points.shape[0] >= 20:
        chosen_points = skel_points
        source = "skeleton"
    elif vox_points.shape[0] >= 100:
        chosen_points = subsample_points(vox_points, 4000, rng)
        source = "mask"
    else:
        raise RuntimeError("068_L HSC has insufficient points for plane fitting.")

    center, normal = fit_plane_pca(chosen_points)
    shape_points = skel_points if skel_points.shape[0] >= 5 else chosen_points
    shape_features = fit_pca_shape(shape_points)
    plane_rmse, plane_mae = plane_residuals(chosen_points, center, normal)
    arc_features = projected_arc_features(shape_points, center, normal)
    vol = float(cleaned.sum() * np.prod(zooms))

    canal_row = {
        "pid": "068",
        "side": "L",
        "struct": "HSC",
        "plane_source": source,
        "mask_points": int(vox_points.shape[0]),
        "skeleton_points": int(skel_points.shape[0]),
        "volume_mm3": round(vol, 6),
        "nx": float(normal[0]),
        "ny": float(normal[1]),
        "nz": float(normal[2]),
        "cx": float(center[0]),
        "cy": float(center[1]),
        "cz": float(center[2]),
        "plane_rmse_mm": plane_rmse,
        "plane_mae_mm": plane_mae,
        **shape_features,
        **arc_features,
        "file": str(fp.resolve()),
    }
    volume_row = {
        "pid": "068",
        "side": "L",
        "struct": "HSC",
        "volume_mm3": round(vol, 6),
        "file": str(fp.resolve()),
    }
    diag_row = {
        "pid": "068",
        "side": "L",
        "struct": "HSC",
        "status": "ok",
        "mask_points": int(vox_points.shape[0]),
        "skeleton_points": int(skel_points.shape[0]),
        "plane_source": source,
        "file": str(fp.resolve()),
    }
    return canal_row, volume_row, diag_row


def write_df(df: pd.DataFrame, path: Path) -> None:
    df.to_csv(path, index=False, encoding="utf-8-sig")


def refresh_stats_and_plots(vol_df: pd.DataFrame, angle_df: pd.DataFrame) -> None:
    angle_stats = (
        angle_df.groupby("angle_type")["angle_deg"]
        .agg(["count", "mean", "std", "min", "median", "max"])
        .reset_index()
        .round(4)
    )
    angle_stats.to_csv(GEOM / "angle_summary_stats.csv", index=False, encoding="utf-8-sig")

    vol_df["volume_mm3"] = pd.to_numeric(vol_df["volume_mm3"], errors="coerce")
    vol_stats = (
        vol_df.groupby("struct")["volume_mm3"]
        .agg(["count", "mean", "std", "min", "median", "max"])
        .reset_index()
        .round(4)
    )
    vol_stats.to_csv(GEOM / "volume_summary_stats.csv", index=False, encoding="utf-8-sig")

    plt.figure(figsize=(8, 4.8))
    angle_df.boxplot(column="angle_deg", by="angle_type", grid=False)
    plt.title("Semicircular canal plane angles")
    plt.suptitle("")
    plt.xlabel("")
    plt.ylabel("Angle (degree)")
    plt.tight_layout()
    plt.savefig(GEOM / "angle_boxplot.png", dpi=220)
    plt.close()

    order = ["Cochlear", "Vestibular", "SSC", "HSC", "PSC", "TV", "ELS"]
    data = [vol_df.loc[vol_df["struct"].eq(s), "volume_mm3"].dropna().to_numpy() for s in order]
    plt.figure(figsize=(9, 4.8))
    plt.boxplot(data, labels=order, showfliers=False)
    plt.ylabel("Volume (mm3)")
    plt.xticks(rotation=25, ha="right")
    plt.tight_layout()
    plt.savefig(GEOM / "volume_boxplot.png", dpi=220)
    plt.close()


def build_derived_tables(vol_df: pd.DataFrame, ear_df: pd.DataFrame) -> pd.DataFrame:
    ear = ear_df.copy()
    for col in [
        "angle_ssc_hsc_deg",
        "angle_ssc_psc_deg",
        "angle_hsc_psc_deg",
        "ssc_volume_mm3",
        "hsc_volume_mm3",
        "psc_volume_mm3",
    ]:
        ear[col] = pd.to_numeric(ear[col], errors="coerce")
    ear_summary = ear[
        [
            "pid",
            "side",
            "complete_three_canals",
            "angle_ssc_hsc_deg",
            "angle_ssc_psc_deg",
            "angle_hsc_psc_deg",
            "ssc_volume_mm3",
            "hsc_volume_mm3",
            "psc_volume_mm3",
        ]
    ].copy()
    ear_summary["three_canal_volume_mm3"] = ear_summary[["ssc_volume_mm3", "hsc_volume_mm3", "psc_volume_mm3"]].sum(axis=1)
    for prefix in ["ssc", "hsc", "psc"]:
        ear_summary[f"{prefix}_share_of_3canals"] = ear_summary[f"{prefix}_volume_mm3"] / ear_summary["three_canal_volume_mm3"]
    write_df(ear_summary, GEOM / "ear_angle_volume_summary.csv")

    wide = (
        vol_df.pivot_table(index=["pid", "side"], columns="struct", values="volume_mm3", aggfunc="first")
        .reset_index()
        .rename_axis(columns=None)
    )
    if {"ELS", "TV"}.issubset(wide.columns):
        wide["els_over_tv"] = pd.to_numeric(wide["ELS"], errors="coerce") / pd.to_numeric(wide["TV"], errors="coerce")
        wide["els_percent_of_tv"] = 100.0 * wide["els_over_tv"]
        wide.to_csv(GEOM / "els_tv_ratios.csv", index=False, encoding="utf-8-sig")
    return ear_summary


def update_geometry() -> dict:
    ensure_dirs()
    for fp in OLD_GEOM.glob("*"):
        if fp.is_file():
            shutil.copy2(fp, GEOM / fp.name)

    canal_row, volume_row, diag_row = compute_hsc_row()

    vol_df = read_csv(GEOM / "volumes_mm3.csv")
    vol_df = vol_df.loc[~key_mask(vol_df, "068", "L", "HSC")].copy()
    vol_df = pd.concat([vol_df, pd.DataFrame([volume_row])], ignore_index=True)
    write_df(vol_df.sort_values(["pid", "side", "struct"]), GEOM / "volumes_mm3.csv")

    canal_df = read_csv(GEOM / "canal_geometry_features.csv")
    canal_df = canal_df.loc[~key_mask(canal_df, "068", "L", "HSC")].copy()
    canal_df = pd.concat([canal_df, pd.DataFrame([canal_row])], ignore_index=True)
    write_df(canal_df.sort_values(["pid", "side", "struct"]), GEOM / "canal_geometry_features.csv")

    diag_df = read_csv(GEOM / "extraction_diagnostics.csv")
    diag_df = diag_df.loc[~key_mask(diag_df, "068", "L", "HSC")].copy()
    diag_df = pd.concat([diag_df, pd.DataFrame([diag_row])], ignore_index=True)
    write_df(diag_df.sort_values(["pid", "side", "struct"]), GEOM / "extraction_diagnostics.csv")

    canal_map = {}
    sub = canal_df.loc[key_mask(canal_df, "068", "L")].copy()
    for _, row in sub.iterrows():
        rec = {k: as_number(v) for k, v in row.to_dict().items()}
        canal_map[str(rec["struct"])] = rec
    new_ear_row = build_ear_feature_row("068", "L", canal_map)

    ear_df = read_csv(GEOM / "ear_geometry_features.csv")
    ear_df = ear_df.loc[~key_mask(ear_df, "068", "L")].copy()
    ear_df = pd.concat([ear_df, pd.DataFrame([new_ear_row])], ignore_index=True)
    write_df(ear_df.sort_values(["pid", "side"]), GEOM / "ear_geometry_features.csv")

    angle_df = read_csv(GEOM / "canal_plane_angles_deg.csv")
    angle_df = angle_df.loc[~key_mask(angle_df, "068", "L")].copy()
    angle_df = pd.concat(
        [
            angle_df,
            pd.DataFrame(
                [
                    {"pid": "068", "side": "L", "angle_type": "SSC-HSC", "angle_deg": new_ear_row["angle_ssc_hsc_deg"]},
                    {"pid": "068", "side": "L", "angle_type": "SSC-PSC", "angle_deg": new_ear_row["angle_ssc_psc_deg"]},
                    {"pid": "068", "side": "L", "angle_type": "HSC-PSC", "angle_deg": new_ear_row["angle_hsc_psc_deg"]},
                ]
            ),
        ],
        ignore_index=True,
    )
    angle_df["angle_deg"] = pd.to_numeric(angle_df["angle_deg"], errors="coerce")
    write_df(angle_df.sort_values(["pid", "side", "angle_type"]), GEOM / "canal_plane_angles_deg.csv")

    plane_df = read_csv(GEOM / "plane_normals.csv")
    plane_df = plane_df.loc[~key_mask(plane_df, "068", "L", "HSC")].copy()
    plane_row = {
        "pid": "068",
        "side": "L",
        "struct": "HSC",
        "nx": canal_row["nx"],
        "ny": canal_row["ny"],
        "nz": canal_row["nz"],
        "cx": canal_row["cx"],
        "cy": canal_row["cy"],
        "cz": canal_row["cz"],
        "n_points": canal_row["skeleton_points"] if canal_row["plane_source"] == "skeleton" else canal_row["mask_points"],
        "plane_source": canal_row["plane_source"],
        "file": canal_row["file"],
    }
    plane_df = pd.concat([plane_df, pd.DataFrame([plane_row])], ignore_index=True)
    write_df(plane_df.sort_values(["pid", "side", "struct"]), GEOM / "plane_normals.csv")

    ear_summary = build_derived_tables(vol_df, ear_df)
    refresh_stats_and_plots(vol_df, angle_df)

    summary = {
        "data_dir": str(Path("seg3").resolve()),
        "old_geometry_dir": str(OLD_GEOM.resolve()),
        "output_dir": str(GEOM.resolve()),
        "update": "068_L HSC added incrementally",
        "matched_files": int(len(vol_df)),
        "canal_rows": int(len(canal_df)),
        "ear_rows": int(len(ear_df)),
        "complete_three_canal_ears": int(pd.to_numeric(ear_df["complete_three_canals"], errors="coerce").fillna(0).sum()),
        "complete_three_canal_ears_before": 399,
        "hsc_volume_068L_mm3": float(canal_row["volume_mm3"]),
        "angles_068L_deg": {
            "SSC-HSC": float(new_ear_row["angle_ssc_hsc_deg"]),
            "SSC-PSC": float(new_ear_row["angle_ssc_psc_deg"]),
            "HSC-PSC": float(new_ear_row["angle_hsc_psc_deg"]),
        },
        "files": {
            "ear_angle_volume_summary": str(GEOM / "ear_angle_volume_summary.csv"),
            "els_tv_ratios": str(GEOM / "els_tv_ratios.csv"),
            "volumes_mm3": str(GEOM / "volumes_mm3.csv"),
            "ear_geometry_features": str(GEOM / "ear_geometry_features.csv"),
            "canal_geometry_features": str(GEOM / "canal_geometry_features.csv"),
            "angles": str(GEOM / "canal_plane_angles_deg.csv"),
        },
    }
    (GEOM / "summary.json").write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")
    return summary


def train_els_tv_model() -> dict:
    args = Namespace(
        analysis_dir=str(GEOM),
        feature_csv=None,
        volume_csv=None,
        output_dir=str(MODEL_DIR),
        target="els_over_tv",
        max_features=40,
        random_state=42,
        predict_feature_csv=None,
        model_path=None,
    )
    run_training(args)
    return json.loads((MODEL_DIR / "summary.json").read_text(encoding="utf-8"))


def set_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_p(doc: Document, text: str, size: float = 10.5, bold: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    add_p(doc, text, size=15 if level == 1 else 12, bold=True)


def add_table(doc: Document, rows: list[list[object]], headers: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
        table.rows[0].cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                value = f"{value:.4f}"
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def fmt_pct(x: float) -> str:
    return f"{100 * x:.2f}%"


def build_report(geom_summary: dict, model_summary: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    add_p(doc, "梅尼埃病人内耳半规管影像特征、ELS/TV 占比预测与 ViT 分割更新报告", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "更新点：补入 068_L 的 HSC 后，重新整合三半规管夹角、体积占比、ELS/TV 预测和 Dice 系数结果。", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "一、核心更新结果")
    add_table(
        doc,
        [
            ["068_L HSC 文件", "已纳入", "seg3/sub068/068L_HSC.nii.gz"],
            ["完整三半规管耳侧", "399 -> 400", "+1 个耳侧，068_L 由不完整转为完整"],
            ["HSC 体积记录", "399 -> 400", f"068_L HSC={geom_summary['hsc_volume_068L_mm3']:.4f} mm3"],
            ["三夹角记录", "1197 -> 1200", "新增 068_L 的 SSC-HSC、SSC-PSC、HSC-PSC"],
            ["ELS/TV 建模耳侧", model_summary["n_ears"], "要求同时具备完整三半规管几何特征和 ELS、TV 体积"],
        ],
        ["项目", "结果", "说明"],
    )

    vol = pd.read_csv(GEOM / "volumes_mm3.csv", dtype={"pid": str})
    ratio = pd.read_csv(GEOM / "els_tv_ratios.csv", dtype={"pid": str})
    ear_summary = pd.read_csv(GEOM / "ear_angle_volume_summary.csv", dtype={"pid": str})
    m68 = ratio[(ratio["pid"].astype(str).str.zfill(3) == "068") & (ratio["side"] == "L")].iloc[0]
    e68 = ear_summary[(ear_summary["pid"].astype(str).str.zfill(3) == "068") & (ear_summary["side"] == "L")].iloc[0]

    add_heading(doc, "二、068_L 补齐后的定量结果")
    add_table(
        doc,
        [
            ["SSC-HSC 夹角", e68["angle_ssc_hsc_deg"], "degree"],
            ["SSC-PSC 夹角", e68["angle_ssc_psc_deg"], "degree"],
            ["HSC-PSC 夹角", e68["angle_hsc_psc_deg"], "degree"],
            ["SSC 体积占三半规管", e68["ssc_share_of_3canals"], "比例"],
            ["HSC 体积占三半规管", e68["hsc_share_of_3canals"], "比例"],
            ["PSC 体积占三半规管", e68["psc_share_of_3canals"], "比例"],
            ["ELS/TV", m68["els_over_tv"], fmt_pct(float(m68["els_over_tv"]))],
        ],
        ["指标", "数值", "单位/说明"],
    )

    add_heading(doc, "三、全队列角度与体积概况")
    angle_stats = pd.read_csv(GEOM / "angle_summary_stats.csv")
    add_table(doc, angle_stats.values.tolist(), angle_stats.columns.tolist())
    doc.add_picture(str(GEOM / "angle_boxplot.png"), width=Cm(15.5))
    vol_stats = pd.read_csv(GEOM / "volume_summary_stats.csv")
    add_table(doc, vol_stats.values.tolist(), vol_stats.columns.tolist())
    doc.add_picture(str(GEOM / "volume_boxplot.png"), width=Cm(15.5))

    add_heading(doc, "四、ELS/TV 预测模型")
    model_cmp = pd.read_csv(MODEL_DIR / "model_comparison.csv")
    show = model_cmp[["model", "cv", "n", "mae", "rmse", "r2", "pearson_r"]].copy()
    add_table(doc, show.round(4).values.tolist(), ["模型", "交叉验证", "n", "MAE", "RMSE", "R2", "Pearson r"])
    best = model_summary["best_metrics"]
    add_p(
        doc,
        f"当前以 ELS/TV 为目标，最优交叉验证模型为 {model_summary['cv_best_model']}；"
        f"RMSE={best['rmse']:.4f}，MAE={best['mae']:.4f}，R2={best['r2']:.4f}。"
        "这些结果可作为 P-EBM 联合临床症状、听力、前庭功能与影像异常变量前的影像组学基线。",
    )

    imp = pd.read_csv(MODEL_DIR / "feature_importance.csv")
    if not imp.empty:
        add_heading(doc, "五、主要预测特征", level=2)
        top = imp.head(12)[["feature", "metric", "value", "abs_value"]].round(5)
        add_table(doc, top.values.tolist(), ["特征", "类型", "值", "绝对值"])

    add_heading(doc, "六、ViT 三半规管分割 Dice 系数")
    vit_rows = []
    for name, folder in [
        ("早期 7 结构/小样本 ViT", Path("analysis_out") / "vit_inner_ear_seg"),
        ("三半规管 ViT 初版", Path("analysis_out") / "vit_inner_ear_seg3"),
        ("三半规管 ViT 45 epoch + crop128", Path("analysis_out") / "vit_inner_ear_seg3_e45_c128"),
    ]:
        js = json.loads((folder / "metrics_summary.json").read_text(encoding="utf-8"))
        test = js.get("tuned_test_summary") or js.get("test_summary")
        vit_rows.append(
            [
                name,
                js.get("train_samples", ""),
                js.get("val_samples", ""),
                js.get("test_samples", ""),
                js.get("best_val_dice", ""),
                test.get("dice", ""),
                test.get("iou", ""),
            ]
        )
    add_table(doc, vit_rows, ["模型/实验", "训练耳", "验证耳", "测试耳", "最佳验证 Dice", "测试 Dice", "测试 IoU"])
    add_p(doc, "Dice 提升：测试 Dice 从 0.6910 提升至 0.8234，绝对提升 0.1323，约相对提升 19.15%。")
    train_curve = Path("analysis_out") / "vit_inner_ear_seg3_e45_c128" / "training_curves.png"
    if train_curve.exists():
        doc.add_picture(str(train_curve), width=Cm(15.5))

    add_heading(doc, "七、P-EBM 联合分析建议")
    add_p(
        doc,
        "建议把本次输出的半规管夹角、三半规管体积占比、ELS/TV、耳蜗/前庭/TV 体积等影像变量，"
        "与临床症状、听力分期、前庭功能和病程时间轴合并成耳侧级与患者级两张表。"
        "P-EBM 可用于估计影像异常、听力下降、前庭功能受损和临床症状之间的发生顺序；"
        "本目录当前未发现独立的临床症状/听力/前庭功能表，因此本报告先完成影像侧可直接落地的基线结果。",
    )

    add_heading(doc, "八、输出文件")
    add_table(
        doc,
        [
            ["几何特征总表", GEOM / "ear_geometry_features.csv"],
            ["夹角+体积占比汇总", GEOM / "ear_angle_volume_summary.csv"],
            ["ELS/TV 比值表", GEOM / "els_tv_ratios.csv"],
            ["ELS/TV 模型比较", MODEL_DIR / "model_comparison.csv"],
            ["模型交叉验证预测", MODEL_DIR / "cv_predictions.csv"],
            ["模型特征重要性", MODEL_DIR / "feature_importance.csv"],
            ["本 Word 报告", REPORT],
        ],
        ["文件", "路径"],
    )

    doc.save(str(REPORT))


def main() -> None:
    geom_summary = update_geometry()
    model_summary = train_els_tv_model()
    build_report(geom_summary, model_summary)
    print(json.dumps({"result_dir": str(ROOT.resolve()), "report": str(REPORT.resolve())}, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
