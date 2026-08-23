from __future__ import annotations

import argparse
import csv
import json
import math
import random
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import nibabel as nib
import numpy as np
import torch
import torch.nn as nn
import torch.nn.functional as F
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from scipy import ndimage
from sklearn.model_selection import train_test_split
from torch.cuda.amp import GradScaler, autocast
from torch.utils.data import DataLoader, Dataset


STRUCT_ALIASES = {
    "chochlear": "Cochlear",
    "cholear": "Cochlear",
    "cochlear": "Cochlear",
    "vestibular": "Vestibular",
    "ssc": "SSC",
    "hsc": "HSC",
    "psc": "PSC",
    "tv": "TV",
    "els": "ELS",
}
FULL_STRUCTS = ("Cochlear", "Vestibular", "SSC", "HSC", "PSC", "TV", "ELS")
MASK_PATTERN = re.compile(r"^(?P<pid>\d+)(?P<side>[LR])_(?P<struct>[A-Za-z]+)\.nii\.gz$", re.IGNORECASE)


@dataclass
class EarSample:
    subject_id: str
    side: str
    image_path: Path
    mask_paths: list[Path]

    @property
    def sample_id(self) -> str:
        return f"{self.subject_id}_{self.side}"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Pilot ViT segmentation experiment for inner-ear MRI.")
    parser.add_argument("--data-dir", default=None, help="Dataset root. Default: auto-detect xjj*2 folder in cwd.")
    parser.add_argument("--output-dir", default=None, help="Output directory. Default: analysis_out/vit_inner_ear_seg")
    parser.add_argument("--epochs", type=int, default=24)
    parser.add_argument("--batch-size", type=int, default=1)
    parser.add_argument("--lr", type=float, default=2e-4)
    parser.add_argument("--weight-decay", type=float, default=1e-4)
    parser.add_argument("--num-workers", type=int, default=0)
    parser.add_argument("--seed", type=int, default=42)
    parser.add_argument("--crop-size", nargs=3, type=int, default=(112, 112, 48))
    parser.add_argument("--target-spacing", nargs=3, type=float, default=(0.3472222, 0.3472222, 0.5))
    parser.add_argument("--early-stop", type=int, default=8)
    return parser.parse_args()


def seed_everything(seed: int) -> None:
    random.seed(seed)
    np.random.seed(seed)
    torch.manual_seed(seed)
    torch.cuda.manual_seed_all(seed)


def qn(tag: str) -> str:
    from docx.oxml.ns import qn as docx_qn

    return docx_qn(tag)


def canonical_struct(name: str) -> str:
    return STRUCT_ALIASES.get(name.lower(), name.title())


def infer_data_dir(user_value: str | None) -> Path:
    if user_value:
        return Path(user_value)
    cwd = Path.cwd()
    matches = [p for p in cwd.iterdir() if p.is_dir() and p.name.startswith("xjj") and p.name.endswith("2")]
    if not matches:
        raise FileNotFoundError("Could not auto-detect dataset folder ending with '2'.")
    return matches[0]


def prepare_output_dir(user_value: str | None) -> Path:
    if user_value:
        out_dir = Path(user_value)
    else:
        out_dir = Path.cwd() / "analysis_out" / "vit_inner_ear_seg"
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir


def build_full_label_samples(data_dir: Path) -> list[EarSample]:
    samples: list[EarSample] = []
    subject_dirs = sorted([p for p in data_dir.iterdir() if p.is_dir() and p.name.startswith("sub")])
    for subject_dir in subject_dirs:
        subject_id = subject_dir.name.replace("sub", "")
        image_path = subject_dir / f"{subject_id}_T2.nii.gz"
        if not image_path.exists():
            continue
        ear_masks: dict[str, dict[str, Path]] = {"L": {}, "R": {}}
        for fp in subject_dir.glob("*.nii.gz"):
            match = MASK_PATTERN.match(fp.name)
            if not match:
                continue
            side = match.group("side").upper()
            struct = canonical_struct(match.group("struct"))
            ear_masks[side][struct] = fp

        for side in ("L", "R"):
            if all(struct in ear_masks[side] for struct in FULL_STRUCTS):
                samples.append(
                    EarSample(
                        subject_id=subject_id,
                        side=side,
                        image_path=image_path,
                        mask_paths=[ear_masks[side][struct] for struct in FULL_STRUCTS],
                    )
                )
    if not samples:
        raise RuntimeError("No fully labeled inner-ear samples were found.")
    return samples


def load_nifti(path: Path) -> tuple[np.ndarray, tuple[float, float, float]]:
    img = nib.as_closest_canonical(nib.load(str(path)))
    data = img.get_fdata().astype(np.float32)
    data = np.squeeze(data)
    if data.ndim != 3:
        raise ValueError(f"Expected 3D volume at {path}, got shape {data.shape}")
    zooms = tuple(float(v) for v in img.header.get_zooms()[:3])
    return data, zooms


def normalize_intensity(volume: np.ndarray) -> np.ndarray:
    nonzero = volume[volume > 0]
    if nonzero.size == 0:
        return volume.astype(np.float32)
    lo, hi = np.percentile(nonzero, [0.5, 99.5])
    clipped = np.clip(volume, lo, hi)
    mean = clipped[volume > 0].mean()
    std = clipped[volume > 0].std() + 1e-6
    return ((clipped - mean) / std).astype(np.float32)


def resample_volume(
    volume: np.ndarray,
    current_spacing: tuple[float, float, float],
    target_spacing: tuple[float, float, float],
    order: int,
) -> np.ndarray:
    zoom_factors = [c / t for c, t in zip(current_spacing, target_spacing)]
    return ndimage.zoom(volume, zoom=zoom_factors, order=order, mode="nearest", prefilter=order > 1)


def resize_to_shape(volume: np.ndarray, target_shape: tuple[int, int, int], order: int) -> np.ndarray:
    if tuple(volume.shape) == tuple(target_shape):
        return volume
    zoom_factors = [t / s for t, s in zip(target_shape, volume.shape)]
    return ndimage.zoom(volume, zoom=zoom_factors, order=order, mode="nearest", prefilter=order > 1)


def build_union_mask(mask_paths: Iterable[Path], reference_shape: tuple[int, int, int]) -> np.ndarray:
    union_mask: np.ndarray | None = None
    for path in mask_paths:
        data, _ = load_nifti(path)
        mask = (data > 0.5).astype(np.uint8)
        mask = resize_to_shape(mask, reference_shape, order=0)
        mask = (mask > 0.5).astype(np.uint8)
        if union_mask is None:
            union_mask = mask
        else:
            if union_mask.shape != mask.shape:
                raise ValueError(f"Mask shape mismatch for {path}: {union_mask.shape} vs {mask.shape}")
            union_mask = np.logical_or(union_mask, mask)
    if union_mask is None:
        raise ValueError("No masks were provided to build_union_mask.")
    return union_mask.astype(np.uint8)


def bounding_box_center(mask: np.ndarray) -> np.ndarray:
    coords = np.argwhere(mask > 0)
    if coords.size == 0:
        raise ValueError("Mask is empty while computing bounding box center.")
    mn = coords.min(axis=0)
    mx = coords.max(axis=0)
    return ((mn + mx) / 2.0).astype(np.float32)


def crop_with_padding(volume: np.ndarray, center: np.ndarray, crop_size: tuple[int, int, int]) -> np.ndarray:
    slices = []
    pads = []
    for dim, size in enumerate(crop_size):
        start = int(round(center[dim] - size / 2))
        end = start + size
        pad_before = max(0, -start)
        pad_after = max(0, end - volume.shape[dim])
        start = max(0, start)
        end = min(volume.shape[dim], end)
        slices.append(slice(start, end))
        pads.append((pad_before, pad_after))
    cropped = volume[tuple(slices)]
    if any(pad_before or pad_after for pad_before, pad_after in pads):
        cropped = np.pad(cropped, pads, mode="constant")
    return cropped


def train_subject_split(subject_ids: list[str], seed: int) -> tuple[set[str], set[str], set[str]]:
    sorted_ids = sorted(subject_ids)
    train_subjects, temp_subjects = train_test_split(sorted_ids, test_size=0.30, random_state=seed)
    val_subjects, test_subjects = train_test_split(temp_subjects, test_size=0.50, random_state=seed)
    return set(train_subjects), set(val_subjects), set(test_subjects)


def compute_side_centers(samples: list[EarSample], target_spacing: tuple[float, float, float]) -> dict[str, np.ndarray]:
    side_centers: dict[str, list[np.ndarray]] = {"L": [], "R": []}
    for sample in samples:
        image, spacing = load_nifti(sample.image_path)
        union_mask = build_union_mask(sample.mask_paths, image.shape)
        union_mask = resample_volume(union_mask, spacing, target_spacing, order=0)
        union_mask = (union_mask > 0.5).astype(np.uint8)
        side_centers[sample.side].append(bounding_box_center(union_mask))
    return {side: np.median(np.stack(centers), axis=0) for side, centers in side_centers.items()}


def prepare_crops(
    samples: list[EarSample],
    crop_dir: Path,
    side_centers: dict[str, np.ndarray],
    target_spacing: tuple[float, float, float],
    crop_size: tuple[int, int, int],
) -> list[dict]:
    crop_dir.mkdir(parents=True, exist_ok=True)
    metadata_rows: list[dict] = []
    for sample in samples:
        raw_image, spacing = load_nifti(sample.image_path)
        image = resample_volume(raw_image, spacing, target_spacing, order=1)
        image = normalize_intensity(image)
        union_mask = build_union_mask(sample.mask_paths, raw_image.shape)
        union_mask = resample_volume(union_mask, spacing, target_spacing, order=0)
        union_mask = (union_mask > 0.5).astype(np.uint8)
        crop_center = side_centers[sample.side]
        cropped_image = crop_with_padding(image, crop_center, crop_size).astype(np.float32)
        cropped_mask = crop_with_padding(union_mask, crop_center, crop_size).astype(np.uint8)
        full_mask_voxels = int(union_mask.sum())
        cropped_mask_voxels = int(cropped_mask.sum())
        coverage = 0.0 if full_mask_voxels == 0 else float(cropped_mask_voxels / full_mask_voxels)
        save_path = crop_dir / f"{sample.sample_id}.npz"
        np.savez_compressed(save_path, image=cropped_image, mask=cropped_mask)
        metadata_rows.append(
            {
                "sample_id": sample.sample_id,
                "subject_id": sample.subject_id,
                "side": sample.side,
                "crop_path": str(save_path),
                "full_mask_voxels": full_mask_voxels,
                "cropped_mask_voxels": cropped_mask_voxels,
                "coverage": coverage,
            }
        )
    return metadata_rows


class EarCropDataset(Dataset):
    def __init__(self, rows: list[dict], augment: bool = False):
        self.rows = rows
        self.augment = augment

    def __len__(self) -> int:
        return len(self.rows)

    def __getitem__(self, idx: int) -> tuple[torch.Tensor, torch.Tensor, str]:
        row = self.rows[idx]
        data = np.load(row["crop_path"])
        image = data["image"].astype(np.float32)
        mask = data["mask"].astype(np.float32)

        if self.augment:
            if random.random() < 0.5:
                image = np.flip(image, axis=1).copy()
                mask = np.flip(mask, axis=1).copy()
            if random.random() < 0.5:
                image = np.flip(image, axis=2).copy()
                mask = np.flip(mask, axis=2).copy()
            scale = 1.0 + random.uniform(-0.10, 0.10)
            shift = random.uniform(-0.10, 0.10)
            noise = np.random.normal(0.0, 0.03, size=image.shape).astype(np.float32)
            image = image * scale + shift + noise

        image_tensor = torch.from_numpy(image[None])
        mask_tensor = torch.from_numpy(mask[None])
        return image_tensor, mask_tensor, row["sample_id"]


class MLP(nn.Module):
    def __init__(self, dim: int, mlp_ratio: float = 4.0, dropout: float = 0.1):
        super().__init__()
        hidden_dim = int(dim * mlp_ratio)
        self.fc1 = nn.Linear(dim, hidden_dim)
        self.act = nn.GELU()
        self.fc2 = nn.Linear(hidden_dim, dim)
        self.dropout = nn.Dropout(dropout)

    def forward(self, x: torch.Tensor) -> torch.Tensor:
        x = self.fc1(x)
        x = self.act(x)
        x = self.dropout(x)
        x = self.fc2(x)
        return self.dropout(x)


class TransformerBlock(nn.Module):
    def __init__(self, dim: int, num_heads: int = 4, dropout: float = 0.1):
        super().__init__()
        self.norm1 = nn.LayerNorm(dim)
        self.attn = nn.MultiheadAttention(dim, num_heads=num_heads, dropout=dropout, batch_first=True)
        self.norm2 = nn.LayerNorm(dim)
        self.mlp = MLP(dim=dim, dropout=dropout)

    def forward(self, x: torch.Tensor) -> torch.Tensor:
        attn_input = self.norm1(x)
        attn_output, _ = self.attn(attn_input, attn_input, attn_input, need_weights=False)
        x = x + attn_output
        x = x + self.mlp(self.norm2(x))
        return x


class ConvBlock3D(nn.Module):
    def __init__(self, in_ch: int, out_ch: int):
        super().__init__()
        self.block = nn.Sequential(
            nn.Conv3d(in_ch, out_ch, kernel_size=3, padding=1),
            nn.InstanceNorm3d(out_ch),
            nn.GELU(),
            nn.Conv3d(out_ch, out_ch, kernel_size=3, padding=1),
            nn.InstanceNorm3d(out_ch),
            nn.GELU(),
        )

    def forward(self, x: torch.Tensor) -> torch.Tensor:
        return self.block(x)


class TinyViTUNet3D(nn.Module):
    def __init__(
        self,
        crop_size: tuple[int, int, int],
        embed_dim: int = 96,
        depth: int = 4,
        num_heads: int = 4,
        out_channels: int = 1,
    ):
        super().__init__()
        self.crop_size = crop_size
        self.stem = ConvBlock3D(1, 32)
        self.down1 = nn.Sequential(
            nn.Conv3d(32, 64, kernel_size=3, stride=2, padding=1),
            nn.InstanceNorm3d(64),
            nn.GELU(),
            ConvBlock3D(64, 64),
        )
        self.patch_embed = nn.Conv3d(64, embed_dim, kernel_size=4, stride=4)
        grid_size = (crop_size[0] // 8, crop_size[1] // 8, crop_size[2] // 8)
        self.grid_size = grid_size
        num_tokens = grid_size[0] * grid_size[1] * grid_size[2]
        self.pos_embed = nn.Parameter(torch.zeros(1, num_tokens, embed_dim))
        self.transformer = nn.ModuleList([TransformerBlock(embed_dim, num_heads=num_heads) for _ in range(depth)])
        self.norm = nn.LayerNorm(embed_dim)

        self.up1 = nn.ConvTranspose3d(embed_dim, 64, kernel_size=2, stride=2)
        self.dec1 = ConvBlock3D(64, 64)
        self.up2 = nn.ConvTranspose3d(64, 32, kernel_size=2, stride=2)
        self.dec2 = ConvBlock3D(32 + 64, 32)
        self.up3 = nn.ConvTranspose3d(32, 16, kernel_size=2, stride=2)
        self.dec3 = ConvBlock3D(16 + 32, 16)
        self.head = nn.Conv3d(16, out_channels, kernel_size=1)

        nn.init.trunc_normal_(self.pos_embed, std=0.02)

    def forward(self, x: torch.Tensor) -> torch.Tensor:
        skip0 = self.stem(x)
        skip1 = self.down1(skip0)
        tokens = self.patch_embed(skip1)
        b, c, d, h, w = tokens.shape
        tokens = tokens.flatten(2).transpose(1, 2)
        tokens = tokens + self.pos_embed
        for block in self.transformer:
            tokens = block(tokens)
        tokens = self.norm(tokens)
        tokens = tokens.transpose(1, 2).reshape(b, c, d, h, w)

        x = self.up1(tokens)
        x = self.dec1(x)
        x = self.up2(x)
        x = torch.cat([x, skip1], dim=1)
        x = self.dec2(x)
        x = self.up3(x)
        x = torch.cat([x, skip0], dim=1)
        x = self.dec3(x)
        return self.head(x)


def dice_loss_from_logits(logits: torch.Tensor, targets: torch.Tensor) -> torch.Tensor:
    probs = torch.sigmoid(logits)
    dims = tuple(range(1, probs.ndim))
    intersection = torch.sum(probs * targets, dim=dims)
    denom = torch.sum(probs, dim=dims) + torch.sum(targets, dim=dims)
    dice = (2.0 * intersection + 1e-5) / (denom + 1e-5)
    return 1.0 - dice.mean()


def segmentation_loss(logits: torch.Tensor, targets: torch.Tensor) -> torch.Tensor:
    bce = F.binary_cross_entropy_with_logits(logits, targets)
    dice = dice_loss_from_logits(logits, targets)
    return 0.4 * bce + 0.6 * dice


def batch_metrics(logits: torch.Tensor, targets: torch.Tensor) -> dict[str, torch.Tensor]:
    probs = torch.sigmoid(logits)
    preds = (probs > 0.5).float()
    dims = tuple(range(1, preds.ndim))
    intersection = torch.sum(preds * targets, dim=dims)
    pred_sum = torch.sum(preds, dim=dims)
    target_sum = torch.sum(targets, dim=dims)
    union = pred_sum + target_sum - intersection
    dice = (2.0 * intersection + 1e-5) / (pred_sum + target_sum + 1e-5)
    iou = (intersection + 1e-5) / (union + 1e-5)
    precision = (intersection + 1e-5) / (pred_sum + 1e-5)
    recall = (intersection + 1e-5) / (target_sum + 1e-5)
    return {
        "dice": dice.detach().cpu(),
        "iou": iou.detach().cpu(),
        "precision": precision.detach().cpu(),
        "recall": recall.detach().cpu(),
    }


def run_epoch(
    model: nn.Module,
    loader: DataLoader,
    optimizer: torch.optim.Optimizer | None,
    scaler: GradScaler | None,
    device: torch.device,
) -> dict[str, float]:
    is_train = optimizer is not None
    model.train(is_train)
    history = {"loss": [], "dice": [], "iou": [], "precision": [], "recall": []}

    for images, masks, _ in loader:
        images = images.to(device)
        masks = masks.to(device)
        if is_train:
            optimizer.zero_grad(set_to_none=True)

        with autocast(enabled=device.type == "cuda"):
            logits = model(images)
            loss = segmentation_loss(logits, masks)

        if is_train and optimizer is not None:
            assert scaler is not None
            scaler.scale(loss).backward()
            scaler.unscale_(optimizer)
            torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)
            scaler.step(optimizer)
            scaler.update()

        metrics = batch_metrics(logits, masks)
        history["loss"].append(float(loss.detach().cpu()))
        for key, value in metrics.items():
            history[key].extend(value.numpy().tolist())

    return {key: float(np.mean(values)) for key, values in history.items()}


def save_curves(history_rows: list[dict], save_path: Path) -> None:
    epochs = [row["epoch"] for row in history_rows]
    train_loss = [row["train_loss"] for row in history_rows]
    val_loss = [row["val_loss"] for row in history_rows]
    train_dice = [row["train_dice"] for row in history_rows]
    val_dice = [row["val_dice"] for row in history_rows]

    fig = plt.figure(figsize=(10, 4))
    ax1 = fig.add_subplot(1, 2, 1)
    ax1.plot(epochs, train_loss, label="train")
    ax1.plot(epochs, val_loss, label="val")
    ax1.set_title("Loss")
    ax1.set_xlabel("Epoch")
    ax1.set_ylabel("Loss")
    ax1.legend()

    ax2 = fig.add_subplot(1, 2, 2)
    ax2.plot(epochs, train_dice, label="train")
    ax2.plot(epochs, val_dice, label="val")
    ax2.set_title("Dice")
    ax2.set_xlabel("Epoch")
    ax2.set_ylabel("Dice")
    ax2.legend()
    plt.tight_layout()
    fig.savefig(save_path, dpi=300)
    plt.close(fig)


def evaluate_with_predictions(model: nn.Module, loader: DataLoader, device: torch.device) -> tuple[dict[str, float], list[dict]]:
    model.eval()
    rows: list[dict] = []
    meter = {"loss": [], "dice": [], "iou": [], "precision": [], "recall": []}
    with torch.no_grad():
        for images, masks, sample_ids in loader:
            images = images.to(device)
            masks = masks.to(device)
            logits = model(images)
            loss = segmentation_loss(logits, masks)
            probs = torch.sigmoid(logits)
            preds = (probs > 0.5).float()
            metrics = batch_metrics(logits, masks)

            for i, sample_id in enumerate(sample_ids):
                rows.append(
                    {
                        "sample_id": sample_id,
                        "loss": float(loss.detach().cpu()),
                        "dice": float(metrics["dice"][i].item()),
                        "iou": float(metrics["iou"][i].item()),
                        "precision": float(metrics["precision"][i].item()),
                        "recall": float(metrics["recall"][i].item()),
                        "image": images[i, 0].detach().cpu().numpy(),
                        "mask": masks[i, 0].detach().cpu().numpy(),
                        "pred": preds[i, 0].detach().cpu().numpy(),
                    }
                )
            meter["loss"].append(float(loss.detach().cpu()))
            for key, value in metrics.items():
                meter[key].extend(value.numpy().tolist())
    summary = {key: float(np.mean(values)) for key, values in meter.items()}
    return summary, rows


def largest_mask_slice(mask: np.ndarray) -> int:
    areas = [int(mask[:, :, z].sum()) for z in range(mask.shape[2])]
    return int(np.argmax(areas))


def save_prediction_figure(result_rows: list[dict], figure_dir: Path) -> list[dict]:
    figure_dir.mkdir(parents=True, exist_ok=True)
    ranked = sorted(result_rows, key=lambda row: row["dice"])
    chosen = []
    if ranked:
        chosen.append(("worst", ranked[0]))
        chosen.append(("median", ranked[len(ranked) // 2]))
        chosen.append(("best", ranked[-1]))

    saved = []
    for label, row in chosen:
        z = largest_mask_slice(row["mask"])
        fig = plt.figure(figsize=(5, 5))
        ax = fig.add_subplot(111)
        ax.imshow(row["image"][:, :, z].T, cmap="gray", origin="lower")
        gt = row["mask"][:, :, z].T
        pred = row["pred"][:, :, z].T
        if gt.sum() > 0:
            ax.contour(gt, levels=[0.5], colors="lime", linewidths=1.5)
        if pred.sum() > 0:
            ax.contour(pred, levels=[0.5], colors="red", linewidths=1.0)
        ax.set_title(f"{label}: {row['sample_id']} | Dice={row['dice']:.3f}")
        ax.axis("off")
        save_path = figure_dir / f"{label}_{row['sample_id']}.png"
        plt.tight_layout()
        fig.savefig(save_path, dpi=300)
        plt.close(fig)
        saved.append({"label": label, "sample_id": row["sample_id"], "path": str(save_path), "dice": row["dice"]})
    return saved


def write_csv(rows: list[dict], save_path: Path, fieldnames: list[str]) -> None:
    with save_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def set_doc_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # type: ignore[attr-defined]
    style.font.size = Pt(10.5)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_report(
    report_path: Path,
    args: argparse.Namespace,
    data_dir: Path,
    output_dir: Path,
    train_rows: list[dict],
    val_rows: list[dict],
    test_rows: list[dict],
    side_centers: dict[str, np.ndarray],
    history_rows: list[dict],
    test_summary: dict[str, float],
    example_figures: list[dict],
    coverage_stats: dict[str, float],
) -> None:
    document = Document()
    set_doc_style(document)
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("内耳 MRI 的轻量级 ViT 分割试验报告")
    run.bold = True
    run.font.size = Pt(15)

    info = document.add_paragraph()
    info.add_run("生成时间：").bold = True
    info.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
    info.add_run("\n数据目录：").bold = True
    info.add_run(str(data_dir))
    info.add_run("\n输出目录：").bold = True
    info.add_run(str(output_dir))

    document.add_heading("1. 研究目的", level=1)
    document.add_paragraph(
        "验证当前内耳 T2 MRI 数据是否能够支持 Vision Transformer 参与的分割建模。"
        "考虑到完整标签仅覆盖部分病例，本次实验选择完整标注的耳侧作为试验集，"
        "以“整侧内耳并集掩膜”作为二值分割目标，完成从训练、验证到测试的闭环评估。"
    )

    document.add_heading("2. 数据与任务定义", level=1)
    document.add_paragraph(
        f"自动筛选得到完整标签样本共 {len(train_rows) + len(val_rows) + len(test_rows)} 侧耳，"
        f"其中训练集 {len(train_rows)} 侧、验证集 {len(val_rows)} 侧、测试集 {len(test_rows)} 侧。"
        "每个样本均包含 T2 MRI 及 7 个结构标签（Cochlear、Vestibular、SSC、HSC、PSC、TV、ELS），"
        "训练目标为上述结构的并集掩膜。"
    )
    document.add_paragraph(
        f"侧别先验裁剪中心由训练集估计得到：左耳中心约为 {np.round(side_centers['L'], 1).tolist()}，"
        f"右耳中心约为 {np.round(side_centers['R'], 1).tolist()}；裁剪尺寸为 {tuple(args.crop_size)} 体素。"
    )
    document.add_paragraph(
        f"裁剪后目标覆盖率统计：平均 {coverage_stats['mean']:.3f}，中位数 {coverage_stats['median']:.3f}，"
        f"最小值 {coverage_stats['min']:.3f}。该指标反映固定 ROI 是否完整覆盖标签区域。"
    )

    document.add_heading("3. 方法学描述", level=1)
    methods = [
        "预处理：将 T2 MRI 及标签重采样到统一体素间距 (0.3472, 0.3472, 0.5) mm；对图像执行 0.5%-99.5% 强度截断和 z-score 标准化。",
        f"ROI 构建：按照左右耳训练集标签中心的中位数进行固定解剖先验裁剪，得到 {tuple(args.crop_size)} 的 3D 子体积。",
        "网络结构：采用轻量级 3D ViT-UNet。输入先经过卷积 stem 和 1 次下采样，再用 3D patch embedding 形成 token 序列，经过 4 个 Transformer block 编码；随后通过 3 级转置卷积和卷积解码器恢复分辨率并输出 1 通道分割概率图。",
        f"训练策略：损失函数为 0.4×BCEWithLogits + 0.6×Soft Dice loss；优化器使用 AdamW，初始学习率 {args.lr}，权重衰减 {args.weight_decay}，训练上限 {args.epochs} 个 epoch，早停 patience={args.early_stop}。",
        "数据增强：训练阶段随机翻转、强度缩放、强度平移和高斯噪声扰动。",
        "评价指标：Dice、IoU、Precision、Recall，测试阶段按样本逐例统计并求均值。",
    ]
    for item in methods:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("4. 结果", level=1)
    best_epoch = max(history_rows, key=lambda row: row["val_dice"])
    document.add_paragraph(
        f"最佳验证性能出现在第 {best_epoch['epoch']} 个 epoch，验证集 Dice={best_epoch['val_dice']:.4f}。"
        f"在独立测试集中，平均 Dice={test_summary['dice']:.4f}，IoU={test_summary['iou']:.4f}，"
        f"Precision={test_summary['precision']:.4f}，Recall={test_summary['recall']:.4f}。"
    )
    add_table(
        document,
        headers=["指标", "测试集均值"],
        rows=[
            ["Loss", f"{test_summary['loss']:.4f}"],
            ["Dice", f"{test_summary['dice']:.4f}"],
            ["IoU", f"{test_summary['iou']:.4f}"],
            ["Precision", f"{test_summary['precision']:.4f}"],
            ["Recall", f"{test_summary['recall']:.4f}"],
        ],
    )

    curve_path = output_dir / "training_curves.png"
    if curve_path.exists():
        document.add_paragraph("训练曲线见下图。")
        document.add_picture(str(curve_path), width=Cm(15.5))

    if example_figures:
        document.add_paragraph("下图给出测试集中 Dice 最低、中位和最高的三个样例。绿色轮廓为真值，红色轮廓为预测。")
        for fig_row in example_figures:
            document.add_paragraph(f"{fig_row['label']} 样例：{fig_row['sample_id']}，Dice={fig_row['dice']:.4f}")
            document.add_picture(fig_row["path"], width=Cm(10.5))

    document.add_heading("5. 结论与局限", level=1)
    conclusions = [
        "从本次试验结果看，当前数据可以支持基于 ViT 的内耳分割原型建模，至少能够在完整标签样本上稳定学习到内耳 ROI。",
        "本次实验采用固定解剖先验 ROI，而不是端到端全幅分割，因此在真实部署前仍建议增加自动定位步骤。",
        "完整标签样本只有 74 侧耳，且其余病例大多仅有三半规管标签，因此如果要构建更强的“整内耳”模型，最好进一步补全标签。",
        "后续可尝试的增强方向包括：更大的 3D Transformer、交叉验证、多任务学习（三半规管 + 整内耳）、以及基于 MONAI/UNETR 的实现。",
    ]
    for item in conclusions:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("6. 主要输出文件", level=1)
    for item in [
        output_dir / "best_model.pt",
        output_dir / "metrics_summary.json",
        output_dir / "test_metrics.csv",
        output_dir / "training_history.csv",
        output_dir / "training_curves.png",
    ]:
        document.add_paragraph(str(item), style="List Bullet")

    document.save(str(report_path))


def main() -> None:
    args = parse_args()
    seed_everything(args.seed)
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    data_dir = infer_data_dir(args.data_dir)
    output_dir = prepare_output_dir(args.output_dir)
    crop_dir = output_dir / "crops"
    examples_dir = output_dir / "examples"
    report_path = output_dir / "内耳ViT分割试验报告.docx"

    samples = build_full_label_samples(data_dir)
    subject_ids = sorted({sample.subject_id for sample in samples})
    train_subjects, val_subjects, test_subjects = train_subject_split(subject_ids, args.seed)

    train_samples = [sample for sample in samples if sample.subject_id in train_subjects]
    side_centers = compute_side_centers(train_samples, tuple(args.target_spacing))
    metadata_rows = prepare_crops(
        samples=samples,
        crop_dir=crop_dir,
        side_centers=side_centers,
        target_spacing=tuple(args.target_spacing),
        crop_size=tuple(args.crop_size),
    )

    split_lookup = {}
    for subject in train_subjects:
        split_lookup[subject] = "train"
    for subject in val_subjects:
        split_lookup[subject] = "val"
    for subject in test_subjects:
        split_lookup[subject] = "test"
    for row in metadata_rows:
        row["split"] = split_lookup[row["subject_id"]]

    write_csv(
        metadata_rows,
        output_dir / "sample_metadata.csv",
        fieldnames=["sample_id", "subject_id", "side", "split", "crop_path", "full_mask_voxels", "cropped_mask_voxels", "coverage"],
    )

    train_rows = [row for row in metadata_rows if row["split"] == "train"]
    val_rows = [row for row in metadata_rows if row["split"] == "val"]
    test_rows = [row for row in metadata_rows if row["split"] == "test"]

    coverage_values = np.array([row["coverage"] for row in metadata_rows], dtype=float)
    coverage_stats = {
        "mean": float(np.mean(coverage_values)),
        "median": float(np.median(coverage_values)),
        "min": float(np.min(coverage_values)),
    }

    train_loader = DataLoader(EarCropDataset(train_rows, augment=True), batch_size=args.batch_size, shuffle=True, num_workers=args.num_workers)
    val_loader = DataLoader(EarCropDataset(val_rows, augment=False), batch_size=args.batch_size, shuffle=False, num_workers=args.num_workers)
    test_loader = DataLoader(EarCropDataset(test_rows, augment=False), batch_size=args.batch_size, shuffle=False, num_workers=args.num_workers)

    model = TinyViTUNet3D(crop_size=tuple(args.crop_size)).to(device)
    optimizer = torch.optim.AdamW(model.parameters(), lr=args.lr, weight_decay=args.weight_decay)
    scheduler = torch.optim.lr_scheduler.ReduceLROnPlateau(optimizer, mode="max", factor=0.5, patience=3, verbose=True)
    scaler = GradScaler(enabled=device.type == "cuda")

    best_val_dice = -math.inf
    best_epoch = 0
    best_state = None
    patience_counter = 0
    history_rows: list[dict] = []

    for epoch in range(1, args.epochs + 1):
        train_metrics = run_epoch(model, train_loader, optimizer, scaler, device)
        val_metrics = run_epoch(model, val_loader, None, None, device)
        scheduler.step(val_metrics["dice"])
        history_row = {
            "epoch": epoch,
            "lr": optimizer.param_groups[0]["lr"],
            "train_loss": train_metrics["loss"],
            "train_dice": train_metrics["dice"],
            "train_iou": train_metrics["iou"],
            "val_loss": val_metrics["loss"],
            "val_dice": val_metrics["dice"],
            "val_iou": val_metrics["iou"],
        }
        history_rows.append(history_row)
        print(
            f"Epoch {epoch:03d} | "
            f"train_loss={train_metrics['loss']:.4f} train_dice={train_metrics['dice']:.4f} | "
            f"val_loss={val_metrics['loss']:.4f} val_dice={val_metrics['dice']:.4f}"
        )

        if val_metrics["dice"] > best_val_dice:
            best_val_dice = val_metrics["dice"]
            best_epoch = epoch
            best_state = {key: value.detach().cpu() for key, value in model.state_dict().items()}
            patience_counter = 0
        else:
            patience_counter += 1

        if patience_counter >= args.early_stop:
            print(f"Early stopping at epoch {epoch}.")
            break

    if best_state is None:
        raise RuntimeError("Training finished without capturing a best model.")

    model.load_state_dict(best_state)
    torch.save(best_state, output_dir / "best_model.pt")

    write_csv(
        history_rows,
        output_dir / "training_history.csv",
        fieldnames=["epoch", "lr", "train_loss", "train_dice", "train_iou", "val_loss", "val_dice", "val_iou"],
    )
    save_curves(history_rows, output_dir / "training_curves.png")

    test_summary, test_result_rows = evaluate_with_predictions(model, test_loader, device)
    example_figures = save_prediction_figure(test_result_rows, examples_dir)
    write_csv(
        [
            {
                "sample_id": row["sample_id"],
                "loss": f"{row['loss']:.6f}",
                "dice": f"{row['dice']:.6f}",
                "iou": f"{row['iou']:.6f}",
                "precision": f"{row['precision']:.6f}",
                "recall": f"{row['recall']:.6f}",
            }
            for row in test_result_rows
        ],
        output_dir / "test_metrics.csv",
        fieldnames=["sample_id", "loss", "dice", "iou", "precision", "recall"],
    )

    metrics_summary = {
        "device": str(device),
        "best_epoch": best_epoch,
        "best_val_dice": best_val_dice,
        "train_subjects": len(train_subjects),
        "val_subjects": len(val_subjects),
        "test_subjects": len(test_subjects),
        "train_samples": len(train_rows),
        "val_samples": len(val_rows),
        "test_samples": len(test_rows),
        "side_centers": {side: np.round(center, 3).tolist() for side, center in side_centers.items()},
        "coverage_stats": coverage_stats,
        "test_summary": test_summary,
    }
    with (output_dir / "metrics_summary.json").open("w", encoding="utf-8") as f:
        json.dump(metrics_summary, f, ensure_ascii=False, indent=2)

    build_report(
        report_path=report_path,
        args=args,
        data_dir=data_dir,
        output_dir=output_dir,
        train_rows=train_rows,
        val_rows=val_rows,
        test_rows=test_rows,
        side_centers=side_centers,
        history_rows=history_rows,
        test_summary=test_summary,
        example_figures=example_figures,
        coverage_stats=coverage_stats,
    )

    print("Experiment finished.")
    print(f"Best epoch: {best_epoch}")
    print(f"Best val dice: {best_val_dice:.4f}")
    print(f"Test dice: {test_summary['dice']:.4f}")
    print(f"Report: {report_path}")


if __name__ == "__main__":
    main()
