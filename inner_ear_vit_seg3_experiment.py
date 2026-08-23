from __future__ import annotations

import argparse
import json
import math
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import numpy as np
import nibabel as nib
import torch
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from torch.cuda.amp import GradScaler
from torch.utils.data import DataLoader

from inner_ear_vit_seg_experiment import (
    EarCropDataset,
    TinyViTUNet3D,
    bounding_box_center,
    build_union_mask,
    compute_side_centers,
    crop_with_padding,
    load_nifti,
    normalize_intensity,
    prepare_crops,
    resample_volume,
    run_epoch,
    save_curves,
    save_prediction_figure,
    segmentation_loss,
    train_subject_split,
    write_csv,
)


STRUCT_ALIASES = {
    "chochlear": "Cochlear",
    "cholear": "Cochlear",
    "cochlear": "Cochlear",
    "vestibular": "Vestibular",
    "ssc": "SSC",
    "hsc": "HSC",
    "psc": "PSC",
    "tv": "TV",
    "els": "ELS",
}
FULL_STRUCTS = ("Cochlear", "Vestibular", "SSC", "HSC", "PSC", "TV", "ELS")
CANAL_STRUCTS = ("SSC", "HSC", "PSC")
MASK_PATTERN = re.compile(r"^(?:sub)?(?P<pid>\d+)(?P<side>[LR])[_-](?P<struct>[A-Za-z]+)\.nii(?:\.gz)?$", re.IGNORECASE)


@dataclass
class Seg3EarSample:
    subject_id: str
    side: str
    image_path: Path
    mask_paths: list[Path]

    @property
    def sample_id(self) -> str:
        return f"{self.subject_id}_{self.side}"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Seg3 ViT segmentation experiment.")
    parser.add_argument("--data-dir", default="seg3")
    parser.add_argument("--output-dir", default=str(Path("analysis_out") / "vit_inner_ear_seg3"))
    parser.add_argument("--structures", default="SSC,HSC,PSC", help="Comma-separated target masks to merge.")
    parser.add_argument("--epochs", type=int, default=18)
    parser.add_argument("--batch-size", type=int, default=1)
    parser.add_argument("--lr", type=float, default=2e-4)
    parser.add_argument("--weight-decay", type=float, default=1e-4)
    parser.add_argument("--num-workers", type=int, default=0)
    parser.add_argument("--seed", type=int, default=42)
    parser.add_argument("--crop-size", nargs=3, type=int, default=(112, 112, 48))
    parser.add_argument("--target-spacing", nargs=3, type=float, default=(0.3472222, 0.3472222, 0.5))
    parser.add_argument("--early-stop", type=int, default=6)
    parser.add_argument("--threshold-min", type=float, default=0.50)
    parser.add_argument("--threshold-max", type=float, default=0.95)
    parser.add_argument("--threshold-step", type=float, default=0.05)
    parser.add_argument(
        "--min-train-coverage",
        type=float,
        default=0.0,
        help="Drop training crops whose target-mask coverage is below this value. Validation/test rows are not filtered.",
    )
    parser.add_argument(
        "--train-crop-center",
        choices=("fixed", "mask"),
        default="fixed",
        help="Use fixed side centers for all crops, or mask centers for training crops only.",
    )
    parser.add_argument(
        "--crop-center-scope",
        choices=("train", "all"),
        default="train",
        help="Apply adaptive mask/fixed ROI selection to training crops only or to all splits.",
    )
    return parser.parse_args()


def canonical_struct(name: str) -> str:
    return STRUCT_ALIASES.get(name.lower(), name.title())


def find_t2(subject_dir: Path, subject_id: str) -> Path | None:
    candidates = [
        subject_dir / f"{subject_id}_T2.nii.gz",
        subject_dir / f"{subject_id}_T2.nii",
        subject_dir / f"sub{subject_id}_T2.nii.gz",
        subject_dir / f"sub{subject_id}_T2.nii",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    matches = sorted(subject_dir.glob("*_T2.nii*"))
    return matches[0] if matches else None


def is_readable_nifti(path: Path) -> tuple[bool, str]:
    try:
        img = nib.load(str(path))
        _ = img.shape
        return True, ""
    except Exception as exc:  # noqa: BLE001 - report data-quality failures without stopping the full run.
        return False, f"{type(exc).__name__}: {exc}"


def scan_dataset(data_dir: Path, target_structs: tuple[str, ...]) -> tuple[list[Seg3EarSample], dict]:
    samples: list[Seg3EarSample] = []
    subject_dirs = sorted([p for p in data_dir.iterdir() if p.is_dir() and p.name.startswith("sub")])
    audit = {
        "subjects_total": len(subject_dirs),
        "subjects_with_t2": 0,
        "target_ear_samples": 0,
        "target_subjects_with_any_ear": 0,
        "target_subjects_with_both_ears": 0,
        "full7_ear_samples": 0,
        "unreadable_ear_samples": 0,
        "unreadable_files": [],
        "missing_or_incomplete_examples": [],
    }

    for subject_dir in subject_dirs:
        subject_id = subject_dir.name.replace("sub", "")
        image_path = find_t2(subject_dir, subject_id)
        if image_path is not None:
            audit["subjects_with_t2"] += 1

        ear_masks: dict[str, dict[str, Path]] = {"L": {}, "R": {}}
        for fp in subject_dir.glob("*.nii*"):
            match = MASK_PATTERN.match(fp.name)
            if not match:
                continue
            side = match.group("side").upper()
            struct = canonical_struct(match.group("struct"))
            ear_masks[side][struct] = fp

        subject_target_count = 0
        for side in ("L", "R"):
            if all(struct in ear_masks[side] for struct in FULL_STRUCTS):
                audit["full7_ear_samples"] += 1
            if image_path is not None and all(struct in ear_masks[side] for struct in target_structs):
                required_paths = [image_path] + [ear_masks[side][struct] for struct in target_structs]
                bad_files = []
                for required_path in required_paths:
                    ok, reason = is_readable_nifti(required_path)
                    if not ok:
                        bad_files.append({"path": str(required_path), "reason": reason})
                if bad_files:
                    audit["unreadable_ear_samples"] += 1
                    audit["unreadable_files"].append({"sample_id": f"{subject_id}_{side}", "files": bad_files})
                    continue
                samples.append(
                    Seg3EarSample(
                        subject_id=subject_id,
                        side=side,
                        image_path=image_path,
                        mask_paths=[ear_masks[side][struct] for struct in target_structs],
                    )
                )
                subject_target_count += 1

        audit["target_ear_samples"] += subject_target_count
        if subject_target_count > 0:
            audit["target_subjects_with_any_ear"] += 1
        if subject_target_count == 2:
            audit["target_subjects_with_both_ears"] += 1
        if image_path is None or subject_target_count == 0:
            audit["missing_or_incomplete_examples"].append(
                {
                    "subject_id": subject_id,
                    "has_t2": image_path is not None,
                    "left_structs": sorted(ear_masks["L"].keys()),
                    "right_structs": sorted(ear_masks["R"].keys()),
                }
            )

    if not samples:
        raise RuntimeError(f"No samples found for target structures: {target_structs}")
    return samples, audit


def set_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(document: Document, text: str, size: float = 10.5, bold: bool = False, align=None):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    return paragraph


def add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)


def load_previous_summary(root: Path) -> dict | None:
    path = root / "analysis_out" / "vit_inner_ear_seg" / "metrics_summary.json"
    if not path.exists():
        return None
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def metrics_from_probs(probs: torch.Tensor, targets: torch.Tensor, threshold: float) -> dict[str, torch.Tensor]:
    preds = (probs > threshold).float()
    dims = tuple(range(1, preds.ndim))
    intersection = torch.sum(preds * targets, dim=dims)
    pred_sum = torch.sum(preds, dim=dims)
    target_sum = torch.sum(targets, dim=dims)
    union = pred_sum + target_sum - intersection
    dice = (2.0 * intersection + 1e-5) / (pred_sum + target_sum + 1e-5)
    iou = (intersection + 1e-5) / (union + 1e-5)
    precision = (intersection + 1e-5) / (pred_sum + 1e-5)
    recall = (intersection + 1e-5) / (target_sum + 1e-5)
    return {
        "dice": dice.detach().cpu(),
        "iou": iou.detach().cpu(),
        "precision": precision.detach().cpu(),
        "recall": recall.detach().cpu(),
    }


def evaluate_with_threshold(model: torch.nn.Module, loader: DataLoader, device: torch.device, threshold: float) -> tuple[dict[str, float], list[dict]]:
    model.eval()
    rows: list[dict] = []
    meter = {"loss": [], "dice": [], "iou": [], "precision": [], "recall": []}
    with torch.no_grad():
        for images, masks, sample_ids in loader:
            images = images.to(device)
            masks = masks.to(device)
            logits = model(images)
            loss = segmentation_loss(logits, masks)
            probs = torch.sigmoid(logits)
            preds = (probs > threshold).float()
            metrics = metrics_from_probs(probs, masks, threshold)

            for i, sample_id in enumerate(sample_ids):
                rows.append(
                    {
                        "sample_id": sample_id,
                        "loss": float(loss.detach().cpu()),
                        "dice": float(metrics["dice"][i].item()),
                        "iou": float(metrics["iou"][i].item()),
                        "precision": float(metrics["precision"][i].item()),
                        "recall": float(metrics["recall"][i].item()),
                        "image": images[i, 0].detach().cpu().numpy(),
                        "mask": masks[i, 0].detach().cpu().numpy(),
                        "pred": preds[i, 0].detach().cpu().numpy(),
                    }
                )
            meter["loss"].append(float(loss.detach().cpu()))
            for key, value in metrics.items():
                meter[key].extend(value.numpy().tolist())
    summary = {key: float(np.mean(values)) for key, values in meter.items()}
    return summary, rows


def threshold_values(args: argparse.Namespace) -> list[float]:
    values = []
    current = args.threshold_min
    while current <= args.threshold_max + 1e-9:
        values.append(round(float(current), 4))
        current += args.threshold_step
    return values


def tune_threshold(model: torch.nn.Module, loader: DataLoader, device: torch.device, thresholds: list[float]) -> tuple[float, list[dict]]:
    rows = []
    for threshold in thresholds:
        summary, _ = evaluate_with_threshold(model, loader, device, threshold)
        rows.append(
            {
                "threshold": threshold,
                "dice": summary["dice"],
                "iou": summary["iou"],
                "precision": summary["precision"],
                "recall": summary["recall"],
            }
        )
    best = max(rows, key=lambda row: row["dice"])
    return float(best["threshold"]), rows


def write_report(
    report_path: Path,
    args: argparse.Namespace,
    data_dir: Path,
    output_dir: Path,
    target_structs: tuple[str, ...],
    audit: dict,
    train_rows: list[dict],
    val_rows: list[dict],
    test_rows: list[dict],
    side_centers: dict[str, np.ndarray],
    coverage_stats: dict[str, float],
    history_rows: list[dict],
    default_threshold: float,
    default_test_summary: dict[str, float],
    tuned_threshold: float,
    tuned_test_summary: dict[str, float],
    threshold_rows: list[dict],
    example_figures: list[dict],
    previous_summary: dict | None,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    add_paragraph(document, "seg3 三半规管 ViT 分割试验报告", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, f"数据目录：{data_dir.resolve()}")
    add_paragraph(document, f"输出目录：{output_dir.resolve()}")

    add_heading(document, "1. 试验目的")
    add_paragraph(
        document,
        "本轮以 SSC/HSC/PSC 三半规管为分割目标，不再要求完整 7 结构标签。"
        "检查 seg3 的 200 个病例后发现，大多数病例具备三半规管标签，"
        "因此本次将三半规管掩膜合并为一个二值目标，评估扩大样本后的分割可行性。"
    )
    add_paragraph(
        document,
        "注意：本轮目标是三半规管并集分割，不是完整内耳 7 结构并集分割；报告中的完整 7 结构数量仅作为数据核查信息。"
    )

    add_heading(document, "2. 数据核查")
    add_table(
        document,
        ["项目", "数量"],
        [
            ["seg3 病例目录", audit["subjects_total"]],
            ["存在 T2 图像的病例", audit["subjects_with_t2"]],
            [f"满足 {'+'.join(target_structs)} 的耳侧样本", audit["target_ear_samples"]],
            ["至少 1 侧满足目标标签的病例", audit["target_subjects_with_any_ear"]],
            ["双侧均满足目标标签的病例", audit["target_subjects_with_both_ears"]],
            ["满足完整 7 结构标签的耳侧样本", audit["full7_ear_samples"]],
            ["因文件不可读剔除的耳侧样本", audit["unreadable_ear_samples"]],
        ],
    )
    add_paragraph(
        document,
        f"按病例级随机划分后，训练集 {len(train_rows)} 耳侧、验证集 {len(val_rows)} 耳侧、测试集 {len(test_rows)} 耳侧。"
    )
    add_paragraph(
        document,
        f"ROI 固定裁剪中心：左耳 {np.round(side_centers['L'], 1).tolist()}，右耳 {np.round(side_centers['R'], 1).tolist()}；"
        f"裁剪尺寸 {tuple(args.crop_size)}，重采样体素间距 {tuple(args.target_spacing)} mm。"
    )
    add_paragraph(
        document,
        f"裁剪后标签覆盖率：平均 {coverage_stats['mean']:.3f}，中位数 {coverage_stats['median']:.3f}，最低 {coverage_stats['min']:.3f}。"
    )

    add_heading(document, "3. 方法")
    for item in [
        f"标签目标：将每侧 {', '.join(target_structs)} 三个 mask 合并为一个二值三半规管 ROI。",
        "预处理：T2 图像和 mask 统一重采样；图像按非零体素进行 0.5%-99.5% 截断和 z-score 标准化。",
        "网络：轻量级 3D ViT-UNet，卷积 stem 提取局部纹理，Transformer block 建模 3D token 的长程关系，转置卷积解码恢复到 ROI 分辨率。",
        f"训练：AdamW，学习率 {args.lr}，权重衰减 {args.weight_decay}，损失为 0.4 BCEWithLogits + 0.6 Soft Dice，最多 {args.epochs} 个 epoch，早停 patience={args.early_stop}。",
        "评估：在独立测试集上逐耳侧计算 Dice、IoU、Precision 和 Recall；同时在验证集上扫描阈值，并将最佳验证阈值应用于测试集。",
    ]:
        document.add_paragraph(item, style="List Bullet")

    add_heading(document, "4. 结果")
    best_epoch = max(history_rows, key=lambda row: row["val_dice"])
    add_table(
        document,
        ["指标", "数值"],
        [
            ["最佳 epoch", best_epoch["epoch"]],
            ["最佳验证 Dice", f"{best_epoch['val_dice']:.4f}"],
            [f"默认阈值 {default_threshold:.2f} 测试 Dice", f"{default_test_summary['dice']:.4f}"],
            [f"验证集调优阈值 {tuned_threshold:.2f} 测试 Dice", f"{tuned_test_summary['dice']:.4f}"],
            ["调优阈值测试 IoU", f"{tuned_test_summary['iou']:.4f}"],
            ["调优阈值测试 Precision", f"{tuned_test_summary['precision']:.4f}"],
            ["调优阈值测试 Recall", f"{tuned_test_summary['recall']:.4f}"],
        ],
    )
    add_table(
        document,
        ["阈值", "验证 Dice", "验证 IoU", "验证 Precision", "验证 Recall"],
        [
            [
                f"{row['threshold']:.2f}",
                f"{row['dice']:.4f}",
                f"{row['iou']:.4f}",
                f"{row['precision']:.4f}",
                f"{row['recall']:.4f}",
            ]
            for row in sorted(threshold_rows, key=lambda row: row["dice"], reverse=True)[:6]
        ],
    )
    if previous_summary:
        add_paragraph(
            document,
            f"本轮三半规管目标进入 {audit['target_ear_samples']} 耳侧；完整 7 结构标签不作为本次分割任务的筛选条件。"
        )
    curve_path = output_dir / "training_curves.png"
    if curve_path.exists():
        add_paragraph(document, "训练曲线：")
        document.add_picture(str(curve_path), width=Cm(15.5))
    if example_figures:
        add_paragraph(document, "测试样例叠加图：绿色为人工标签，红色为模型预测。")
        for fig in example_figures:
            add_paragraph(document, f"{fig['label']}：{fig['sample_id']}，Dice={fig['dice']:.4f}")
            document.add_picture(fig["path"], width=Cm(9.5))

    add_heading(document, "5. 结论与建议")
    conclusions = [
        f"以三半规管并集为目标时，可用样本达到 {audit['target_ear_samples']} 个耳侧，已经能够更合理地检验 ViT-UNet 的分割学习能力。",
        f"本轮默认阈值测试 Dice={default_test_summary['dice']:.4f}；验证集调优阈值 {tuned_threshold:.2f} 后，测试 Dice={tuned_test_summary['dice']:.4f}、IoU={tuned_test_summary['iou']:.4f}。",
        "后续若要做可部署分割流程，建议增加自动定位/粗分割步骤，减少当前固定 ROI 对解剖位置稳定性的依赖。",
    ]
    for item in conclusions:
        document.add_paragraph(item, style="List Bullet")

    add_heading(document, "6. 输出文件")
    for item in [
        output_dir / "best_model.pt",
        output_dir / "metrics_summary.json",
        output_dir / "test_metrics.csv",
        output_dir / "test_metrics_threshold_0p50.csv",
        output_dir / "threshold_tuning.csv",
        output_dir / "sample_metadata.csv",
        output_dir / "training_history.csv",
        output_dir / "training_curves.png",
    ]:
        document.add_paragraph(str(item.resolve()), style="List Bullet")

    document.save(str(report_path))


def prepare_crops_with_train_center_mode(
    samples: list[Seg3EarSample],
    crop_dir: Path,
    side_centers: dict[str, np.ndarray],
    target_spacing: tuple[float, float, float],
    crop_size: tuple[int, int, int],
    train_subjects: set[str],
    train_crop_center: str,
    crop_center_scope: str,
) -> list[dict]:
    if train_crop_center == "fixed":
        return prepare_crops(samples, crop_dir, side_centers, target_spacing, crop_size)

    crop_dir.mkdir(parents=True, exist_ok=True)
    metadata_rows: list[dict] = []
    for sample in samples:
        raw_image, spacing = load_nifti(sample.image_path)
        image = resample_volume(raw_image, spacing, target_spacing, order=1)
        image = normalize_intensity(image)
        union_mask = build_union_mask(sample.mask_paths, raw_image.shape)
        union_mask = resample_volume(union_mask, spacing, target_spacing, order=0)
        union_mask = (union_mask > 0.5).astype(np.uint8)
        fixed_center = side_centers[sample.side]
        crop_center = fixed_center
        use_adaptive_center = sample.subject_id in train_subjects or crop_center_scope == "all"
        if use_adaptive_center:
            mask_center = bounding_box_center(union_mask)
            fixed_mask = crop_with_padding(union_mask, fixed_center, crop_size).astype(np.uint8)
            adaptive_mask = crop_with_padding(union_mask, mask_center, crop_size).astype(np.uint8)
            crop_center = mask_center if int(adaptive_mask.sum()) >= int(fixed_mask.sum()) else fixed_center
        cropped_image = crop_with_padding(image, crop_center, crop_size).astype(np.float32)
        cropped_mask = crop_with_padding(union_mask, crop_center, crop_size).astype(np.uint8)
        full_mask_voxels = int(union_mask.sum())
        cropped_mask_voxels = int(cropped_mask.sum())
        coverage = 0.0 if full_mask_voxels == 0 else float(cropped_mask_voxels / full_mask_voxels)
        save_path = crop_dir / f"{sample.sample_id}.npz"
        np.savez_compressed(save_path, image=cropped_image, mask=cropped_mask)
        metadata_rows.append(
            {
                "sample_id": sample.sample_id,
                "subject_id": sample.subject_id,
                "side": sample.side,
                "crop_path": str(save_path),
                "full_mask_voxels": full_mask_voxels,
                "cropped_mask_voxels": cropped_mask_voxels,
                "coverage": coverage,
            }
        )
    return metadata_rows


def main() -> None:
    args = parse_args()
    torch.manual_seed(args.seed)
    np.random.seed(args.seed)

    data_dir = Path(args.data_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    crop_dir = output_dir / "crops"
    examples_dir = output_dir / "examples"
    report_path = output_dir / "seg3三半规管ViT分割报告.docx"

    target_structs = tuple(canonical_struct(item.strip()) for item in args.structures.split(",") if item.strip())
    samples, audit = scan_dataset(data_dir, target_structs)
    subject_ids = sorted({sample.subject_id for sample in samples})
    train_subjects, val_subjects, test_subjects = train_subject_split(subject_ids, args.seed)
    train_samples = [sample for sample in samples if sample.subject_id in train_subjects]

    side_centers = compute_side_centers(train_samples, tuple(args.target_spacing))
    metadata_rows = prepare_crops_with_train_center_mode(
        samples,
        crop_dir,
        side_centers,
        tuple(args.target_spacing),
        tuple(args.crop_size),
        train_subjects,
        args.train_crop_center,
        args.crop_center_scope,
    )

    split_lookup = {subject: "train" for subject in train_subjects}
    split_lookup.update({subject: "val" for subject in val_subjects})
    split_lookup.update({subject: "test" for subject in test_subjects})
    for row in metadata_rows:
        row["split"] = split_lookup[row["subject_id"]]

    write_csv(
        metadata_rows,
        output_dir / "sample_metadata.csv",
        fieldnames=["sample_id", "subject_id", "side", "split", "crop_path", "full_mask_voxels", "cropped_mask_voxels", "coverage"],
    )

    raw_train_rows = [row for row in metadata_rows if row["split"] == "train"]
    train_rows = [row for row in raw_train_rows if float(row["coverage"]) >= float(args.min_train_coverage)]
    val_rows = [row for row in metadata_rows if row["split"] == "val"]
    test_rows = [row for row in metadata_rows if row["split"] == "test"]

    coverage_values = np.array([row["coverage"] for row in metadata_rows], dtype=float)
    coverage_stats = {
        "mean": float(np.mean(coverage_values)),
        "median": float(np.median(coverage_values)),
        "min": float(np.min(coverage_values)),
    }

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    train_loader = DataLoader(EarCropDataset(train_rows, augment=True), batch_size=args.batch_size, shuffle=True, num_workers=args.num_workers)
    val_loader = DataLoader(EarCropDataset(val_rows, augment=False), batch_size=args.batch_size, shuffle=False, num_workers=args.num_workers)
    test_loader = DataLoader(EarCropDataset(test_rows, augment=False), batch_size=args.batch_size, shuffle=False, num_workers=args.num_workers)

    model = TinyViTUNet3D(crop_size=tuple(args.crop_size)).to(device)
    optimizer = torch.optim.AdamW(model.parameters(), lr=args.lr, weight_decay=args.weight_decay)
    scheduler = torch.optim.lr_scheduler.ReduceLROnPlateau(optimizer, mode="max", factor=0.5, patience=3, verbose=True)
    scaler = GradScaler(enabled=device.type == "cuda")

    best_val_dice = -math.inf
    best_epoch = 0
    best_state = None
    patience_counter = 0
    history_rows: list[dict] = []

    print(f"Device: {device}")
    print(f"Target structures: {target_structs}")
    print(f"Samples: train={len(train_rows)}, val={len(val_rows)}, test={len(test_rows)}")
    if args.min_train_coverage > 0:
        print(
            f"Filtered train crops by coverage >= {args.min_train_coverage:.3f}: "
            f"kept {len(train_rows)}/{len(raw_train_rows)}",
            flush=True,
        )
    for epoch in range(1, args.epochs + 1):
        train_metrics = run_epoch(model, train_loader, optimizer, scaler, device)
        val_metrics = run_epoch(model, val_loader, None, None, device)
        scheduler.step(val_metrics["dice"])
        history_row = {
            "epoch": epoch,
            "lr": optimizer.param_groups[0]["lr"],
            "train_loss": train_metrics["loss"],
            "train_dice": train_metrics["dice"],
            "train_iou": train_metrics["iou"],
            "val_loss": val_metrics["loss"],
            "val_dice": val_metrics["dice"],
            "val_iou": val_metrics["iou"],
        }
        history_rows.append(history_row)
        print(
            f"Epoch {epoch:03d} | train_loss={train_metrics['loss']:.4f} "
            f"train_dice={train_metrics['dice']:.4f} | val_loss={val_metrics['loss']:.4f} "
            f"val_dice={val_metrics['dice']:.4f}",
            flush=True,
        )

        if val_metrics["dice"] > best_val_dice:
            best_val_dice = val_metrics["dice"]
            best_epoch = epoch
            best_state = {key: value.detach().cpu() for key, value in model.state_dict().items()}
            patience_counter = 0
        else:
            patience_counter += 1
        if patience_counter >= args.early_stop:
            print(f"Early stopping at epoch {epoch}.")
            break

    if best_state is None:
        raise RuntimeError("Training finished without a best model.")

    model.load_state_dict(best_state)
    torch.save(best_state, output_dir / "best_model.pt")
    write_csv(history_rows, output_dir / "training_history.csv", fieldnames=["epoch", "lr", "train_loss", "train_dice", "train_iou", "val_loss", "val_dice", "val_iou"])
    save_curves(history_rows, output_dir / "training_curves.png")

    default_threshold = 0.5
    default_test_summary, default_test_result_rows = evaluate_with_threshold(model, test_loader, device, default_threshold)
    tuned_threshold, threshold_rows = tune_threshold(model, val_loader, device, threshold_values(args))
    tuned_test_summary, tuned_test_result_rows = evaluate_with_threshold(model, test_loader, device, tuned_threshold)
    example_figures = save_prediction_figure(tuned_test_result_rows, examples_dir)
    write_csv(
        [
            {
                "threshold": f"{row['threshold']:.4f}",
                "dice": f"{row['dice']:.6f}",
                "iou": f"{row['iou']:.6f}",
                "precision": f"{row['precision']:.6f}",
                "recall": f"{row['recall']:.6f}",
            }
            for row in threshold_rows
        ],
        output_dir / "threshold_tuning.csv",
        fieldnames=["threshold", "dice", "iou", "precision", "recall"],
    )
    write_csv(
        [
            {
                "sample_id": row["sample_id"],
                "loss": f"{row['loss']:.6f}",
                "dice": f"{row['dice']:.6f}",
                "iou": f"{row['iou']:.6f}",
                "precision": f"{row['precision']:.6f}",
                "recall": f"{row['recall']:.6f}",
            }
            for row in default_test_result_rows
        ],
        output_dir / "test_metrics_threshold_0p50.csv",
        fieldnames=["sample_id", "loss", "dice", "iou", "precision", "recall"],
    )
    write_csv(
        [
            {
                "sample_id": row["sample_id"],
                "loss": f"{row['loss']:.6f}",
                "dice": f"{row['dice']:.6f}",
                "iou": f"{row['iou']:.6f}",
                "precision": f"{row['precision']:.6f}",
                "recall": f"{row['recall']:.6f}",
            }
            for row in tuned_test_result_rows
        ],
        output_dir / "test_metrics.csv",
        fieldnames=["sample_id", "loss", "dice", "iou", "precision", "recall"],
    )

    metrics_summary = {
        "device": str(device),
        "target_structs": target_structs,
        "dataset_audit": audit,
        "best_epoch": best_epoch,
        "best_val_dice": best_val_dice,
        "train_subjects": len(train_subjects),
        "val_subjects": len(val_subjects),
        "test_subjects": len(test_subjects),
        "raw_train_samples": len(raw_train_rows),
        "train_samples": len(train_rows),
        "val_samples": len(val_rows),
        "test_samples": len(test_rows),
        "min_train_coverage": float(args.min_train_coverage),
        "train_crop_center": args.train_crop_center,
        "crop_center_scope": args.crop_center_scope,
        "side_centers": {side: np.round(center, 3).tolist() for side, center in side_centers.items()},
        "coverage_stats": coverage_stats,
        "default_threshold": default_threshold,
        "default_test_summary": default_test_summary,
        "tuned_threshold": tuned_threshold,
        "tuned_test_summary": tuned_test_summary,
        "test_summary": tuned_test_summary,
    }
    with (output_dir / "metrics_summary.json").open("w", encoding="utf-8") as f:
        json.dump(metrics_summary, f, ensure_ascii=False, indent=2)

    write_report(
        report_path=report_path,
        args=args,
        data_dir=data_dir,
        output_dir=output_dir,
        target_structs=target_structs,
        audit=audit,
        train_rows=train_rows,
        val_rows=val_rows,
        test_rows=test_rows,
        side_centers=side_centers,
        coverage_stats=coverage_stats,
        history_rows=history_rows,
        default_threshold=default_threshold,
        default_test_summary=default_test_summary,
        tuned_threshold=tuned_threshold,
        tuned_test_summary=tuned_test_summary,
        threshold_rows=threshold_rows,
        example_figures=example_figures,
        previous_summary=load_previous_summary(Path.cwd()),
    )

    print("Experiment finished.")
    print(f"Best epoch: {best_epoch}")
    print(f"Best val dice: {best_val_dice:.4f}")
    print(f"Default threshold test dice: {default_test_summary['dice']:.4f}")
    print(f"Tuned threshold: {tuned_threshold:.2f}")
    print(f"Tuned threshold test dice: {tuned_test_summary['dice']:.4f}")
    print(f"Report: {report_path}")


if __name__ == "__main__":
    main()
