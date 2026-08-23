from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc: Document, text: str, size: float = 10.5, bold: bool = False, align=None) -> None:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, size=14, bold=True)


def add_table(doc: Document, frame: pd.DataFrame, max_rows: int | None = None) -> None:
    data = frame.copy()
    if max_rows is not None:
        data = data.head(max_rows)
    table = doc.add_table(rows=1, cols=len(data.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, column in enumerate(data.columns):
        table.rows[0].cells[idx].text = str(column)
        table.rows[0].cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for _, row in data.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                text = f"{value:.4f}"
            else:
                text = str(value)
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def main() -> None:
    out = Path("seg3") / "analysis_out_geometry"
    summary = json.loads((out / "summary.json").read_text(encoding="utf-8"))
    angle_summary = pd.read_csv(out / "angle_summary_stats.csv")
    volume_summary = pd.read_csv(out / "volume_summary_stats.csv")
    diag = pd.read_csv(out / "extraction_diagnostics.csv")

    report = out / "seg3_angle_volume_report.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    add_paragraph(doc, "seg3 三半规管夹角与体积计算报告", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "计算目标：SSC/HSC/PSC 三半规管平面夹角，以及清理后 mask 体积。", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"输出目录：{out.resolve()}")

    add_heading(doc, "1. 数据完成情况")
    overview = pd.DataFrame(
        [
            ["匹配到的 mask 文件", summary["matched_files"]],
            ["半规管几何记录", summary["canal_rows"]],
            ["耳侧记录", summary["ear_rows"]],
            ["完整三半规管夹角耳侧", summary["complete_three_canal_ears"]],
            ["骨架点拟合平面", summary["plane_source_counts"].get("skeleton", 0)],
            ["mask 点云回退拟合平面", summary["plane_source_counts"].get("mask", 0)],
            ["skeleton_min_points", summary["skeleton_min_points"]],
            ["fallback_min_mask_points", summary["fallback_min_mask_points"]],
        ],
        columns=["项目", "数量"],
    )
    add_table(doc, overview)
    add_paragraph(
        doc,
        "说明：文件位于 sub### 目录下时，以父目录病例编号为准；因此 sub128/138R_HSC.nii.gz 已按 128_R_HSC 计入。",
    )

    add_heading(doc, "2. 三半规管夹角统计")
    angle_cn = angle_summary.rename(
        columns={
            "angle_type": "夹角",
            "count": "数量",
            "mean": "均值",
            "std": "标准差",
            "min": "最小值",
            "median": "中位数",
            "max": "最大值",
        }
    )
    add_table(doc, angle_cn)
    doc.add_picture(str(out / "angle_boxplot.png"), width=Cm(15))

    add_heading(doc, "3. 体积统计")
    volume_cn = volume_summary.rename(
        columns={
            "struct": "结构",
            "count": "数量",
            "mean": "均值",
            "std": "标准差",
            "min": "最小值",
            "median": "中位数",
            "max": "最大值",
        }
    )
    add_table(doc, volume_cn)
    doc.add_picture(str(out / "volume_boxplot.png"), width=Cm(15))

    add_heading(doc, "4. 未完成或需注意样本")
    problem = diag[diag["status"] != "ok"][
        ["pid", "side", "struct", "status", "mask_points", "skeleton_points", "plane_source", "file"]
    ].copy()
    if problem.empty:
        add_paragraph(doc, "无。")
    else:
        problem = problem.rename(
            columns={
                "pid": "病例",
                "side": "侧别",
                "struct": "结构",
                "status": "状态",
                "mask_points": "mask点数",
                "skeleton_points": "骨架点数",
                "plane_source": "平面来源",
                "file": "文件",
            }
        )
        add_table(doc, problem, max_rows=20)
        add_paragraph(
            doc,
            "当前最终结果中，仅 068_L 缺少 HSC 文件，不能形成完整三夹角；其余低点数 SSC 已使用 mask 点云回退拟合补全。",
        )

    add_heading(doc, "5. 主要输出文件")
    files = pd.DataFrame(
        [
            ["耳侧夹角+三半规管体积汇总", out / "ear_angle_volume_summary.csv"],
            ["全部体积明细", out / "volumes_mm3.csv"],
            ["三半规管夹角明细", out / "canal_plane_angles_deg.csv"],
            ["半规管平面法向量", out / "plane_normals.csv"],
            ["提取诊断", out / "extraction_diagnostics.csv"],
            ["夹角统计", out / "angle_summary_stats.csv"],
            ["体积统计", out / "volume_summary_stats.csv"],
        ],
        columns=["文件", "路径"],
    )
    add_table(doc, files)
    doc.save(str(report))
    print(report)


if __name__ == "__main__":
    main()
