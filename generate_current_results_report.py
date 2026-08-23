from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 10.5, bold: bool = False) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc: Document, text: str, *, size: float = 10.5, bold: bool = False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", size=16 if level == 1 else 13, bold=True)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[object]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def complete_ears(angle_df: pd.DataFrame) -> int:
    grouped = angle_df.groupby(["pid", "side"])["angle_type"].nunique()
    return int((grouped >= 3).sum())


def usable_ears(angle_df: pd.DataFrame, vol_df: pd.DataFrame) -> int:
    complete = angle_df.groupby(["pid", "side"])["angle_type"].nunique().reset_index()
    complete = complete[complete["angle_type"] >= 3][["pid", "side"]]
    wide = vol_df.pivot_table(index=["pid", "side"], columns="struct", values="volume_mm3", aggfunc="first").reset_index()
    ok = wide[(wide["ELS"].notna()) & (wide["TV"].notna()) & (wide["TV"] > 0)][["pid", "side"]]
    return int(len(complete.merge(ok, on=["pid", "side"], how="inner")))


def main() -> None:
    root = Path(__file__).resolve().parent
    analysis_root = root / "analysis_out"
    analysis_root.mkdir(exist_ok=True)
    report_path = analysis_root / "current_results_summary_utf8_2026-04-07.docx"

    folders = {p.name: p for p in root.iterdir() if p.is_dir() and p.name.startswith("xjj")}
    base1 = next(p for name, p in folders.items() if not name.endswith("2")) / "analysis_out"
    base2 = next(p for name, p in folders.items() if name.endswith("2"))
    orig2 = base2 / "analysis_out"
    s20 = base2 / "analysis_out_s20"
    geom = base2 / "analysis_out_geometry"
    geom_model = geom / "geometry_model"
    angle_model1 = base1 / "angle_ratio_model"
    angle_model2 = s20 / "angle_ratio_model_py312"

    base1_model = pd.read_csv(angle_model1 / "model_comparison.csv")
    s20_model = pd.read_csv(angle_model2 / "model_comparison.csv")
    geom_model_df = pd.read_csv(geom_model / "model_comparison.csv")
    diag_df = pd.read_csv(geom / "extraction_diagnostics.csv")
    orig_angle = pd.read_csv(orig2 / "canal_plane_angles_deg.csv")
    s20_angle = pd.read_csv(s20 / "canal_plane_angles_deg.csv")
    geom_angle = pd.read_csv(geom / "canal_plane_angles_deg.csv")
    orig_vol = pd.read_csv(orig2 / "volumes_mm3.csv")
    s20_vol = pd.read_csv(s20 / "volumes_mm3.csv")
    geom_vol = pd.read_csv(geom / "volumes_mm3.csv")
    geom_feat = pd.read_csv(geom / "ear_geometry_features.csv")

    orig_complete = complete_ears(orig_angle)
    s20_complete = complete_ears(s20_angle)
    geom_complete = complete_ears(geom_angle)
    orig_usable = usable_ears(orig_angle, orig_vol)
    s20_usable = usable_ears(s20_angle, s20_vol)
    geom_usable = usable_ears(geom_angle, geom_vol)

    status_counts = diag_df["status"].value_counts(dropna=False)
    source_counts = diag_df["plane_source"].fillna("NaN").value_counts(dropna=False)
    problem_examples = diag_df[diag_df["status"] != "ok"][
        ["pid", "side", "struct", "status", "mask_points", "skeleton_points", "plane_source"]
    ].head(6)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    add_paragraph(
        doc,
        "内耳半规管空间关系与 ELS/TV 预测现有结果总结报告",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "报告目的：总结目前半规管角度提取修复、增强几何特征提取，以及用几何特征预测内淋巴间隙体积占比（ELS/TV）的现有结果。")

    add_heading(doc, "一、项目目标", 1)
    add_paragraph(doc, "当前分析链路为：半规管分割 -> 三半规管空间关系提取 -> 几何特征构建 -> 使用几何特征预测内淋巴间隙体积占比。")
    add_paragraph(doc, "本轮重点完成了两项工作：1）修复 xjj内耳分割2 中角度提取丢样本的问题；2）扩展几何特征，不再只用三组夹角。")

    add_heading(doc, "二、角度提取修复结果", 1)
    add_paragraph(doc, "原始 xjj内耳分割2 的主要问题不是文件缺失，而是骨架点阈值过严。原脚本默认 skeleton_min_points=30，导致大量半规管被判定为骨架点不足而跳过。")
    add_table(
        doc,
        ["结果目录", "角度记录数", "完整三半规管耳朵数", "可用于 ELS/TV 建模的耳朵数", "说明"],
        [
            ["xjj内耳分割2/analysis_out", len(orig_angle), orig_complete, orig_usable, "原始结果，阈值过严"],
            ["xjj内耳分割2/analysis_out_s20", len(s20_angle), s20_complete, s20_usable, "将 skeleton_min_points 调整为 20 后重跑"],
            ["xjj内耳分割2/analysis_out_geometry", len(geom_angle), geom_complete, geom_usable, "在增强几何提取流程下进一步恢复样本"],
        ],
    )
    add_paragraph(doc, f"关键改动后，完整三半规管耳朵数从 {orig_complete} 提升到 {s20_complete}，再到 {geom_complete}；可用于 ELS/TV 建模的耳朵数从 {orig_usable} 提升到 {s20_usable}，再到 {geom_usable}。")

    add_heading(doc, "三、失败诊断与鲁棒性改进", 1)
    add_table(
        doc,
        ["诊断项", "数量"],
        [
            ["status=ok", int(status_counts.get("ok", 0))],
            ["status=insufficient_points", int(status_counts.get("insufficient_points", 0))],
            ["status=missing_file", int(status_counts.get("missing_file", 0))],
            ["plane_source=skeleton", int(source_counts.get("skeleton", 0))],
            ["plane_source=mask", int(source_counts.get("mask", 0))],
        ],
    )
    add_paragraph(doc, "增强版提取脚本新增了回退策略：若骨架点不足但清理后的 mask 体素足够，则改用 mask 点云做 PCA 平面拟合。这一步让少量边界样本不再直接丢失。")
    if not problem_examples.empty:
        add_table(
            doc,
            ["pid", "side", "struct", "status", "mask_points", "skeleton_points", "plane_source"],
            problem_examples.values.tolist(),
        )

    add_heading(doc, "四、增强几何特征内容", 1)
    add_paragraph(doc, "在新的 inner_ear_geometry_features.py 中，除三组夹角外，还提取了以下特征：")
    for text in [
        "半规管平面法向量：nx、ny、nz，以及其绝对值。",
        "半规管中心位置：cx、cy、cz，以及三个半规管中心之间的距离。",
        "半规管体积：SSC/HSC/PSC 的体积（mm^3）。",
        "PCA 形状特征：axis1/axis2/axis3、linearity、planarity、scattering。",
        "平面拟合误差：plane_rmse_mm、plane_mae_mm。",
        "投影弧形特征：proj_radius_mean_mm、proj_arc_span_deg、proj_arc_length_mm、proj_chord_length_mm、proj_arc_chord_ratio。",
        "左右对称性特征：对侧镜像后的 contra_* 特征和 absdiff_* 差值特征。",
    ]:
        add_paragraph(doc, text)
    add_paragraph(doc, f"当前增强特征表共有 {len(geom_feat)} 行耳朵记录；进入建模的数据中共有 70 只耳朵，特征维度为 266。")

    add_heading(doc, "五、ELS/TV 预测结果对比", 1)
    add_paragraph(doc, "为了公平比较，分别记录了三种情况：1）xjj内耳分割 的角度基线；2）xjj内耳分割2 修复后的角度基线；3）xjj内耳分割2 的增强几何模型。")
    add_table(
        doc,
        ["场景", "最佳/对照模型", "n", "MAE", "RMSE", "R²", "Pearson r"],
        [
            ["xjj内耳分割：仅角度", "ridge", 61, "0.0762", "0.1004", "0.0545", "0.2716"],
            ["xjj内耳分割2：仅角度（s20）", "ridge", 62, "0.0732", "0.0953", "-0.0726", "-0.0279"],
            ["xjj内耳分割2：增强几何", "baseline_mean", 70, "0.0764", "0.0967", "0.0000", "NaN"],
        ],
    )
    add_paragraph(doc, "结果说明：在 xjj内耳分割 中，三组夹角对 ELS/TV 有轻微信号（R² 约 0.05）；但在 xjj内耳分割2 中，无论只用角度还是加入增强几何特征，当前都没有稳定超过均值基线。")
    add_paragraph(doc, "这说明当前阶段的主要瓶颈已不再是角度提不出来，而是现有几何特征对 ELS/TV 的预测信号仍偏弱。")

    add_heading(doc, "六、当前结论", 1)
    for text in [
        f"1. 样本恢复是成功的：xjj内耳分割2 的完整三半规管耳朵数已从 {orig_complete} 提升到 {geom_complete}。",
        "2. 默认阈值 30 明显过严，20 更适合当前这批分割数据。",
        "3. 增强几何特征提取已经搭好，后续可以继续叠加更强的结构关系特征。",
        "4. 现阶段单靠半规管几何，尚不能稳定预测 ELS/TV；如果继续做预测，建议并入更多与前庭、耳蜗、TV/ELS 相对位置相关的特征，或考虑配准后群体比较。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "七、主要结果文件", 1)
    add_table(
        doc,
        ["文件类别", "路径"],
        [
            ["修复后角度结果", str(s20 / "canal_plane_angles_deg.csv")],
            ["增强几何特征（耳朵级）", str(geom / "ear_geometry_features.csv")],
            ["增强几何特征（半规管级）", str(geom / "canal_geometry_features.csv")],
            ["提取诊断", str(geom / "extraction_diagnostics.csv")],
            ["角度基线模型结果", str(angle_model2 / "model_comparison.csv")],
            ["增强几何模型结果", str(geom_model / "model_comparison.csv")],
        ],
    )

    fig_candidates = [s20 / "angles_summary.png", s20 / "volumes_summary.png"]
    rep_img = next(iter(sorted(s20.glob("*_3canals_planes_normals.png"))), None)
    if rep_img is not None:
        fig_candidates.append(rep_img)
    valid_figs = [p for p in fig_candidates if p.exists()]
    if valid_figs:
        add_heading(doc, "八、附图", 1)
        for fig in valid_figs:
            add_paragraph(doc, fig.name)
            doc.add_picture(str(fig), width=Cm(14.5))

    doc.save(str(report_path))
    print(report_path)


if __name__ == "__main__":
    main()
