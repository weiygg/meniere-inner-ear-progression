from __future__ import annotations

import argparse
import csv
import math
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from statistics import median

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import nibabel as nib
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from skimage.measure import label, regionprops
from skimage.morphology import ball, binary_closing, skeletonize


STRUCTS_CANALS = ("SSC", "HSC", "PSC")
STRUCTS_VOLUME = ("Cochlear", "Vestibular", "SSC", "HSC", "PSC", "TV", "ELS")
STRUCT_ALIASES = {
    "COCHLEAR": "Cochlear",
    "CHOCHLEAR": "Cochlear",
    "CHOLEAR": "Cochlear",
    "VESTIBULAR": "Vestibular",
    "SSC": "SSC",
    "HSC": "HSC",
    "PSC": "PSC",
    "TV": "TV",
    "ELS": "ELS",
}
FNAME_PATTERN = re.compile(
    r"(?P<pid>\d+)(?P<side>[LR])[-_](?P<struct>[A-Za-z]+)\.nii(\.gz)?$",
    re.IGNORECASE,
)


def canonical_struct(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z]", "", name).upper()
    return STRUCT_ALIASES.get(cleaned, cleaned.title())


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Inner-ear canal angle modeling and Word summary report.")
    parser.add_argument("--data-dir", required=True, help="Input dataset folder, e.g. xjj内耳分割2")
    parser.add_argument("--output-dir", default=None, help="Output folder; default: <data-dir>/analysis_out")
    parser.add_argument("--report-name", default=None, help="Output report filename")
    parser.add_argument("--min-component-voxels", type=int, default=200)
    parser.add_argument("--closing-radius", type=int, default=1)
    # A threshold of 30 was too strict for the second dataset and discarded most canals.
    parser.add_argument("--skeleton-min-points", type=int, default=20)
    parser.add_argument("--plane-size-mm", type=float, default=20.0)
    parser.add_argument("--subsample-points", type=int, default=2500)
    return parser.parse_args()


def load_mask_nii(fp: Path) -> tuple[np.ndarray, tuple[float, float, float], np.ndarray]:
    img = nib.load(str(fp))
    img = nib.as_closest_canonical(img)
    data = img.get_fdata()
    mask = (data > 0.5).astype(np.uint8)
    zooms = img.header.get_zooms()[:3]
    return mask, zooms, img.affine


def keep_largest_component(mask: np.ndarray, min_voxels: int) -> np.ndarray:
    lab = label(mask)
    if lab.max() == 0:
        return mask
    props = sorted(regionprops(lab), key=lambda r: r.area, reverse=True)
    largest = props[0]
    if largest.area < min_voxels:
        return mask
    return (lab == largest.label).astype(np.uint8)


def clean_mask(mask: np.ndarray, closing_radius: int, min_voxels: int) -> np.ndarray:
    if mask.sum() == 0:
        return mask
    selem = ball(closing_radius)
    closed = binary_closing(mask.astype(bool), selem).astype(np.uint8)
    return keep_largest_component(closed, min_voxels=min_voxels)


def volume_mm3(mask: np.ndarray, zooms: tuple[float, float, float]) -> float:
    return float(mask.sum() * float(zooms[0] * zooms[1] * zooms[2]))


def skeleton_points_mm(mask: np.ndarray, zooms: tuple[float, float, float], min_points: int) -> np.ndarray | None:
    if mask.sum() == 0:
        return None
    sk = skeletonize(mask.astype(bool), method="lee")
    pts = np.argwhere(sk > 0)
    if pts.shape[0] < min_points:
        return None
    return pts.astype(float) * np.array(zooms)[None, :]


def fit_plane_pca(points_mm: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    center = points_mm.mean(axis=0)
    x = points_mm - center
    _, _, vt = np.linalg.svd(x, full_matrices=False)
    normal = vt[-1, :]
    normal = normal / (np.linalg.norm(normal) + 1e-12)
    return center, normal


def angle_between_normals(n1: np.ndarray, n2: np.ndarray) -> float:
    n1 = n1 / (np.linalg.norm(n1) + 1e-12)
    n2 = n2 / (np.linalg.norm(n2) + 1e-12)
    cos_value = np.clip(np.abs(np.dot(n1, n2)), -1.0, 1.0)
    return float(np.degrees(np.arccos(cos_value)))


def set_axes_equal(ax) -> None:
    x_limits = ax.get_xlim3d()
    y_limits = ax.get_ylim3d()
    z_limits = ax.get_zlim3d()
    x_range = abs(x_limits[1] - x_limits[0])
    y_range = abs(y_limits[1] - y_limits[0])
    z_range = abs(z_limits[1] - z_limits[0])
    radius = 0.5 * max([x_range, y_range, z_range])
    x_middle = np.mean(x_limits)
    y_middle = np.mean(y_limits)
    z_middle = np.mean(z_limits)
    ax.set_xlim3d([x_middle - radius, x_middle + radius])
    ax.set_ylim3d([y_middle - radius, y_middle + radius])
    ax.set_zlim3d([z_middle - radius, z_middle + radius])


def plot_skeleton_and_plane(
    points_mm: np.ndarray,
    center: np.ndarray,
    normal: np.ndarray,
    title: str,
    save_path: Path,
    plane_size: float,
) -> None:
    fig = plt.figure(figsize=(7, 6))
    ax = fig.add_subplot(111, projection="3d")
    ax.scatter(points_mm[:, 0], points_mm[:, 1], points_mm[:, 2], s=2, alpha=0.8)

    n = normal / (np.linalg.norm(normal) + 1e-12)
    anchor = np.array([1.0, 0.0, 0.0])
    if np.abs(np.dot(anchor, n)) > 0.9:
        anchor = np.array([0.0, 1.0, 0.0])
    u = np.cross(n, anchor)
    u = u / (np.linalg.norm(u) + 1e-12)
    v = np.cross(n, u)
    v = v / (np.linalg.norm(v) + 1e-12)

    grid = np.linspace(-plane_size, plane_size, 10)
    uu, vv = np.meshgrid(grid, grid)
    plane_pts = center[None, None, :] + uu[..., None] * u[None, None, :] + vv[..., None] * v[None, None, :]
    ax.plot_surface(plane_pts[..., 0], plane_pts[..., 1], plane_pts[..., 2], alpha=0.25, linewidth=0)
    ax.quiver(center[0], center[1], center[2], n[0], n[1], n[2], length=plane_size * 0.9, normalize=True)

    ax.set_title(title)
    ax.set_xlabel("x (mm)")
    ax.set_ylabel("y (mm)")
    ax.set_zlabel("z (mm)")
    ax.view_init(elev=20, azim=35)
    set_axes_equal(ax)
    plt.tight_layout()
    fig.savefig(save_path, dpi=300)
    plt.close(fig)


def plot_three_canals_planes(
    canal_data: dict[str, dict[str, np.ndarray]],
    title: str,
    save_path: Path,
    plane_size: float,
    subsample_points: int,
) -> None:
    rng = np.random.default_rng(42)
    fig = plt.figure(figsize=(8, 7))
    ax = fig.add_subplot(111, projection="3d")

    for struct in STRUCTS_CANALS:
        data = canal_data.get(struct)
        if not data:
            continue
        pts = data["pts"]
        if pts.shape[0] > subsample_points:
            idx = rng.choice(pts.shape[0], subsample_points, replace=False)
            pts = pts[idx]
        ax.scatter(pts[:, 0], pts[:, 1], pts[:, 2], s=2, alpha=0.75, label=struct)

        center = data["c"]
        normal = data["n"] / (np.linalg.norm(data["n"]) + 1e-12)
        anchor = np.array([1.0, 0.0, 0.0])
        if np.abs(np.dot(anchor, normal)) > 0.9:
            anchor = np.array([0.0, 1.0, 0.0])
        u = np.cross(normal, anchor)
        u = u / (np.linalg.norm(u) + 1e-12)
        v = np.cross(normal, u)
        v = v / (np.linalg.norm(v) + 1e-12)

        grid = np.linspace(-plane_size, plane_size, 10)
        uu, vv = np.meshgrid(grid, grid)
        plane_pts = center[None, None, :] + uu[..., None] * u[None, None, :] + vv[..., None] * v[None, None, :]
        ax.plot_surface(plane_pts[..., 0], plane_pts[..., 1], plane_pts[..., 2], alpha=0.18, linewidth=0)
        ax.quiver(center[0], center[1], center[2], normal[0], normal[1], normal[2], length=plane_size * 0.9, normalize=True)

    ax.set_title(title)
    ax.set_xlabel("x (mm)")
    ax.set_ylabel("y (mm)")
    ax.set_zlabel("z (mm)")
    ax.legend(loc="upper right")
    ax.view_init(elev=20, azim=35)
    set_axes_equal(ax)
    plt.tight_layout()
    fig.savefig(save_path, dpi=300)
    plt.close(fig)


def write_csv(rows: list[dict], path: Path, fieldnames: list[str]) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def plot_angle_distributions(angle_rows: list[dict], save_path: Path) -> None:
    values = [row["angle_deg"] for row in angle_rows]
    fig = plt.figure(figsize=(10, 4))
    ax1 = fig.add_subplot(1, 2, 1)
    ax1.hist(values, bins=min(20, max(5, len(values))))
    ax1.set_xlabel("Angle (degrees)")
    ax1.set_ylabel("Count")
    ax1.set_title("Angle distribution")

    grouped: dict[str, list[float]] = defaultdict(list)
    for row in angle_rows:
        grouped[row["angle_type"]].append(row["angle_deg"])

    ax2 = fig.add_subplot(1, 2, 2)
    labels = list(grouped.keys())
    data = [grouped[label] for label in labels]
    ax2.boxplot(data, tick_labels=labels, vert=True)
    ax2.set_ylabel("Angle (degrees)")
    ax2.set_title("By angle type")
    plt.tight_layout()
    fig.savefig(save_path, dpi=300)
    plt.close(fig)


def plot_volume_distributions(volume_rows: list[dict], save_path: Path) -> None:
    grouped: dict[str, list[float]] = defaultdict(list)
    for row in volume_rows:
        grouped[row["struct"]].append(row["volume_mm3"])

    fig = plt.figure(figsize=(10, 4))
    ax = fig.add_subplot(111)
    labels = list(grouped.keys())
    data = [grouped[label] for label in labels]
    ax.boxplot(data, tick_labels=labels, vert=True)
    ax.set_ylabel("Volume (mm^3)")
    ax.set_title("Volume distributions by structure")
    plt.xticks(rotation=30, ha="right")
    plt.tight_layout()
    fig.savefig(save_path, dpi=300)
    plt.close(fig)


def parse_files(data_dir: Path) -> list[dict]:
    items = []
    for fp in sorted(data_dir.rglob("*.nii*")):
        match = FNAME_PATTERN.match(fp.name)
        if not match:
            continue
        items.append(
            {
                "pid": match.group("pid"),
                "side": match.group("side").upper(),
                "struct": canonical_struct(match.group("struct")),
                "path": fp,
            }
        )
    return items


def quantile(values: list[float], q: float) -> float:
    if not values:
        return float("nan")
    arr = np.asarray(sorted(values), dtype=float)
    return float(np.quantile(arr, q))


def mean_std_text(values: list[float]) -> str:
    if not values:
        return "-"
    arr = np.asarray(values, dtype=float)
    mean_value = float(arr.mean())
    std_value = float(arr.std(ddof=1)) if arr.size > 1 else float("nan")
    if math.isnan(std_value):
        return f"{mean_value:.2f}"
    return f"{mean_value:.2f} +/- {std_value:.2f}"


def median_iqr_text(values: list[float]) -> str:
    if not values:
        return "-"
    q1 = quantile(values, 0.25)
    med = quantile(values, 0.50)
    q3 = quantile(values, 0.75)
    return f"{med:.2f} ({q1:.2f}, {q3:.2f})"


def ensure_page_fields(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    fld_start = OxmlElement("w:fldChar")
    fld_start.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_start)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)
    paragraph.add_run(" 页")


def set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def generate_report(
    report_path: Path,
    data_dir: Path,
    output_dir: Path,
    file_rows: list[dict],
    volume_rows: list[dict],
    plane_rows: list[dict],
    angle_rows: list[dict],
) -> None:
    doc = Document()
    set_doc_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{data_dir.name} 内耳三半规管角度计算建模结果总结")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    grouped_ears: dict[tuple[str, str], list[dict]] = defaultdict(list)
    for row in file_rows:
        grouped_ears[(row["pid"], row["side"])].append(row)
    plane_by_ear: dict[tuple[str, str], set[str]] = defaultdict(set)
    for row in plane_rows:
        plane_by_ear[(row["pid"], row["side"])].add(row["struct"])
    angle_by_type: dict[str, list[float]] = defaultdict(list)
    for row in angle_rows:
        angle_by_type[row["angle_type"]].append(row["angle_deg"])
    volume_by_struct: dict[str, list[float]] = defaultdict(list)
    for row in volume_rows:
        volume_by_struct[row["struct"]].append(row["volume_mm3"])

    doc.add_paragraph(f"数据目录：{data_dir}")
    doc.add_paragraph(f"结果目录：{output_dir}")
    doc.add_paragraph(
        "分析流程：递归扫描分割 NIfTI 文件，进行形态学闭运算与最大连通域保留，提取三半规管骨架点云，"
        "采用 PCA/SVD 拟合半规管平面，并基于平面法向量计算夹角。"
    )

    add_heading(doc, "一、总体情况", level=1)
    summary_rows = [
        ["受试者文件夹数", str(len({Path(row['path']).parent.name for row in file_rows}))],
        ["匹配分割文件数", str(len(file_rows))],
        ["耳侧数（pid+side）", str(len(grouped_ears))],
        ["获得平面拟合的耳侧数", str(len(plane_by_ear))],
        ["完整三半规管耳侧数", str(len({(row['pid'], row['side']) for row in angle_rows}))],
        ["角度记录总数", str(len(angle_rows))],
    ]
    add_table(doc, ["指标", "数值"], summary_rows)

    add_heading(doc, "二、三半规管夹角统计", level=1)
    if angle_rows:
        angle_stats_rows = []
        for angle_type in ("SSC-HSC", "SSC-PSC", "HSC-PSC"):
            values = angle_by_type.get(angle_type, [])
            angle_stats_rows.append(
                [
                    angle_type,
                    str(len(values)),
                    mean_std_text(values),
                    median_iqr_text(values),
                    "-" if not values else f"{min(values):.2f}",
                    "-" if not values else f"{max(values):.2f}",
                ]
            )
        add_table(doc, ["角度类型", "n", "均值±SD", "中位数(IQR)", "最小值", "最大值"], angle_stats_rows)

        detail_rows = []
        for row in sorted(angle_rows, key=lambda x: (int(x["pid"]), x["side"], x["angle_type"])):
            detail_rows.append([row["pid"], row["side"], row["angle_type"], f'{row["angle_deg"]:.2f}'])
        add_heading(doc, "三、逐耳角度明细", level=1)
        add_table(doc, ["pid", "侧别", "角度类型", "角度(°)"], detail_rows)
    else:
        doc.add_paragraph("本批次未获得完整三半规管平面组合，因此未生成有效夹角统计。")

    add_heading(doc, "四、体积统计（清理后 mask）", level=1)
    volume_rows_table = []
    for struct in STRUCTS_VOLUME:
        values = volume_by_struct.get(struct, [])
        volume_rows_table.append(
            [
                struct,
                str(len(values)),
                mean_std_text(values),
                median_iqr_text(values),
                "-" if not values else f"{min(values):.2f}",
                "-" if not values else f"{max(values):.2f}",
            ]
        )
    add_table(doc, ["结构", "n", "均值±SD (mm^3)", "中位数(IQR)", "最小值", "最大值"], volume_rows_table)

    add_heading(doc, "五、拟合覆盖情况", level=1)
    coverage_rows = []
    for struct in STRUCTS_CANALS:
        count = sum(1 for structs in plane_by_ear.values() if struct in structs)
        coverage_rows.append([struct, str(count)])
    coverage_rows.append(["三者齐全", str(len({(row['pid'], row['side']) for row in angle_rows}))])
    add_table(doc, ["项目", "耳侧数"], coverage_rows)

    summary_notes = doc.add_paragraph()
    summary_notes.add_run("结果说明：").bold = True
    if angle_rows:
        complete_ears = len({(row["pid"], row["side"]) for row in angle_rows})
        summary_notes.add_run(
            f"本批次共有 {complete_ears} 个耳侧完成三半规管平面拟合并进入角度统计。"
            "原始明细 CSV、法向量 CSV 及可视化 PNG 已同步输出，便于继续核查。"
        )
    else:
        summary_notes.add_run(
            "由于骨架点数量不足或部分半规管缺失，本批次没有耳侧满足完整三半规管夹角统计条件。"
            "但体积统计、单结构拟合图和覆盖情况仍已导出。"
        )

    add_heading(doc, "六、结果图示", level=1)
    angle_plot = output_dir / "angles_summary.png"
    volume_plot = output_dir / "volumes_summary.png"
    if angle_plot.exists():
        doc.add_paragraph("图1  三半规管夹角分布汇总")
        doc.add_picture(str(angle_plot), width=Cm(15.5))
    if volume_plot.exists():
        doc.add_paragraph("图2  各结构体积分布汇总")
        doc.add_picture(str(volume_plot), width=Cm(15.5))

    combined_images = sorted(output_dir.glob("*_3canals_planes_normals.png"))
    if not combined_images:
        combined_images = sorted(output_dir.glob("*_skeleton_plane.png"))
    for idx, img_path in enumerate(combined_images[:4], start=3):
        doc.add_paragraph(f"图{idx}  {img_path.stem}")
        doc.add_picture(str(img_path), width=Cm(14.5))

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    ensure_page_fields(doc.sections[-1].footer.paragraphs[0])
    doc.save(str(report_path))


def main() -> None:
    args = parse_args()
    data_dir = Path(args.data_dir).resolve()
    if not data_dir.exists():
        raise FileNotFoundError(f"Data directory not found: {data_dir}")

    output_dir = Path(args.output_dir).resolve() if args.output_dir else data_dir / "analysis_out"
    output_dir.mkdir(parents=True, exist_ok=True)
    report_name = args.report_name or f"{data_dir.name}_角度计算建模结果总结.docx"
    report_path = output_dir / report_name

    file_rows = parse_files(data_dir)
    if not file_rows:
        raise RuntimeError("未匹配到符合命名规则的 NIfTI 文件。")

    grouped: dict[tuple[str, str], list[dict]] = defaultdict(list)
    for row in file_rows:
        grouped[(row["pid"], row["side"])].append(row)

    volume_rows: list[dict] = []
    plane_rows: list[dict] = []

    for (pid, side), sub_rows in sorted(grouped.items(), key=lambda x: (int(x[0][0]), x[0][1])):
        rows_by_struct: dict[str, list[dict]] = defaultdict(list)
        for row in sub_rows:
            rows_by_struct[row["struct"]].append(row)

        for struct in STRUCTS_VOLUME:
            recs = rows_by_struct.get(struct, [])
            if not recs:
                continue
            fp = recs[0]["path"]
            mask, zooms, _ = load_mask_nii(fp)
            cleaned = clean_mask(mask, closing_radius=args.closing_radius, min_voxels=args.min_component_voxels)
            volume_rows.append(
                {
                    "pid": pid,
                    "side": side,
                    "struct": struct,
                    "volume_mm3": round(volume_mm3(cleaned, zooms), 6),
                    "file": str(fp),
                }
            )

        canal_data: dict[str, dict[str, np.ndarray]] = {}
        for struct in STRUCTS_CANALS:
            recs = rows_by_struct.get(struct, [])
            if not recs:
                continue
            fp = recs[0]["path"]
            mask, zooms, _ = load_mask_nii(fp)
            cleaned = clean_mask(mask, closing_radius=args.closing_radius, min_voxels=args.min_component_voxels)
            points_mm = skeleton_points_mm(cleaned, zooms, min_points=args.skeleton_min_points)
            if points_mm is None:
                continue
            center, normal = fit_plane_pca(points_mm)
            canal_data[struct] = {"pts": points_mm, "c": center, "n": normal}
            plane_rows.append(
                {
                    "pid": pid,
                    "side": side,
                    "struct": struct,
                    "nx": round(float(normal[0]), 8),
                    "ny": round(float(normal[1]), 8),
                    "nz": round(float(normal[2]), 8),
                    "cx": round(float(center[0]), 6),
                    "cy": round(float(center[1]), 6),
                    "cz": round(float(center[2]), 6),
                    "n_points": int(points_mm.shape[0]),
                    "file": str(fp),
                }
            )
            plot_skeleton_and_plane(
                points_mm=points_mm,
                center=center,
                normal=normal,
                title=f"{pid}{side} {struct}: skeleton + plane + normal",
                save_path=output_dir / f"{pid}{side}_{struct}_skeleton_plane.png",
                plane_size=args.plane_size_mm,
            )

        if len(canal_data) >= 2:
            plot_three_canals_planes(
                canal_data=canal_data,
                title=f"{pid}{side} | SSC/HSC/PSC planes + normals",
                save_path=output_dir / f"{pid}{side}_3canals_planes_normals.png",
                plane_size=args.plane_size_mm,
                subsample_points=args.subsample_points,
            )

    angle_rows: list[dict] = []
    grouped_planes: dict[tuple[str, str], dict[str, np.ndarray]] = defaultdict(dict)
    for row in plane_rows:
        grouped_planes[(row["pid"], row["side"])][row["struct"]] = np.array([row["nx"], row["ny"], row["nz"]], dtype=float)

    for (pid, side), normal_map in sorted(grouped_planes.items(), key=lambda x: (int(x[0][0]), x[0][1])):
        if all(struct in normal_map for struct in STRUCTS_CANALS):
            angle_rows.append(
                {"pid": pid, "side": side, "angle_type": "SSC-HSC", "angle_deg": round(angle_between_normals(normal_map["SSC"], normal_map["HSC"]), 6)}
            )
            angle_rows.append(
                {"pid": pid, "side": side, "angle_type": "SSC-PSC", "angle_deg": round(angle_between_normals(normal_map["SSC"], normal_map["PSC"]), 6)}
            )
            angle_rows.append(
                {"pid": pid, "side": side, "angle_type": "HSC-PSC", "angle_deg": round(angle_between_normals(normal_map["HSC"], normal_map["PSC"]), 6)}
            )

    write_csv(volume_rows, output_dir / "volumes_mm3.csv", ["pid", "side", "struct", "volume_mm3", "file"])
    write_csv(
        plane_rows,
        output_dir / "plane_normals.csv",
        ["pid", "side", "struct", "nx", "ny", "nz", "cx", "cy", "cz", "n_points", "file"],
    )
    write_csv(angle_rows, output_dir / "canal_plane_angles_deg.csv", ["pid", "side", "angle_type", "angle_deg"])

    if angle_rows:
        plot_angle_distributions(angle_rows, output_dir / "angles_summary.png")
    if volume_rows:
        plot_volume_distributions(volume_rows, output_dir / "volumes_summary.png")

    generate_report(
        report_path=report_path,
        data_dir=data_dir,
        output_dir=output_dir,
        file_rows=file_rows,
        volume_rows=volume_rows,
        plane_rows=plane_rows,
        angle_rows=angle_rows,
    )

    print(f"Matched files: {len(file_rows)}")
    print(f"Volume rows: {len(volume_rows)}")
    print(f"Plane rows: {len(plane_rows)}")
    print(f"Angle rows: {len(angle_rows)}")
    print(f"Output directory: {output_dir}")
    print(f"Report saved to: {report_path}")


if __name__ == "__main__":
    main()
